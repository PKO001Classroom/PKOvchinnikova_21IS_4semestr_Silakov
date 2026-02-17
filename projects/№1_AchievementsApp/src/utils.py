"""
Вспомогательные функции
"""
import json
import os
from typing import List
from docx import Document
import tkinter.messagebox as messagebox


def load_types_from_json(filepath: str = "types.json") -> List[str]:
    """Загрузка типов достижений из JSON-файла"""
    default_types = ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]
    
    try:
        if os.path.exists(filepath):
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list) and data:
                    return data
        return default_types
    except (json.JSONDecodeError, Exception) as e:
        print(f"Ошибка загрузки {filepath}: {e}")
        return default_types


def export_to_word(achievements: List[tuple], filename: str = "достижения.docx") -> bool:
    """Экспорт достижений в Word-документ"""
    try:
        doc = Document()
        doc.add_heading("Личные учебные достижения", 0)

        if not achievements:
            doc.add_paragraph("Нет сохранённых достижений.")
        else:
            for record in achievements:
                id_num, name, date, typ, level, desc = record

                # Добавляем достижение
                p = doc.add_paragraph()

                # Название - жирным
                title_run = p.add_run(f"{name}")
                title_run.bold = True

                # Дата - курсивом
                date_run = p.add_run(f" — {date}")
                date_run.italic = True

                # Тип и уровень
                p.add_run(f" ({typ}, {level})")

                # Описание (если есть)
                if desc and desc.strip():
                    doc.add_paragraph(f"Описание: {desc}")

                # Разделительная линия
                doc.add_paragraph()

        # Сохраняем документ
        doc.save(filename)
        print(f"Документ сохранён: {filename}")
        return True

    except Exception as e:
        print(f"Ошибка экспорта в Word: {e}")
        return False