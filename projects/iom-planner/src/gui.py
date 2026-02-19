"""
Модуль графического интерфейса
"""
import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import os
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import database
from . import models
from . import utils


class IOMApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Планировщик индивидуального образовательного маршрута")
        self.root.geometry("1200x800")

        # Инициализация БД
        database.init_db()
        database.load_competencies_to_db()

        # Создание виджетов
        self.create_widgets()
        self.check_all_achievements()
        self.update_stats()
        self.update_semester_progress_auto()

    def create_widgets(self):
        """Создание всех вкладок интерфейса"""
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True)

        # Вкладка 1: Мои цели
        self.tab_goals = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_goals, text='Мои цели')
        self.create_goals_tab()

        # Вкладка 2: Мой профиль
        self.tab_profile = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_profile, text='Мой профиль')
        self.create_profile_tab()

        # Вкладка 3: Компетенции
        self.tab_competencies = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_competencies, text='Компетенции')
        self.create_competencies_tab()

        # Вкладка 4: Достижения
        self.tab_achievements = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_achievements, text='Достижения')
        self.create_achievements_tab()

        # Вкладка 5: Цели на семестр
        self.tab_semester = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_semester, text='Цели на семестр')
        self.create_semester_tab()

        # Вкладка 6: Настройки
        self.tab_settings = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_settings, text='Настройки')
        self.create_settings_tab()

    # ============= ВКЛАДКА "МОИ ЦЕЛИ" =============
    def create_goals_tab(self):
        """Создание вкладки со списком целей"""
        list_frame = ttk.Frame(self.tab_goals)
        list_frame.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        # Treeview для отображения целей
        self.goals_tree = ttk.Treeview(list_frame, columns=('ID', 'Название', 'Тип', 'Статус'), show='headings')
        self.goals_tree.heading('ID', text='ID')
        self.goals_tree.heading('Название', text='Название')
        self.goals_tree.heading('Тип', text='Тип')
        self.goals_tree.heading('Статус', text='Статус')
        self.goals_tree.pack(fill='both', expand=True)

        # Кнопки управления
        btn_frame = ttk.Frame(list_frame)
        btn_frame.pack(fill='x', pady=5)

        ttk.Button(btn_frame, text='➕ Добавить', command=self.add_goal).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='✏️ Редактировать', command=self.edit_goal).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='🗑️ Удалить', command=self.delete_goal).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='🔄 Обновить', command=self.refresh_goals).pack(side='left', padx=2)

        self.refresh_goals()

    def refresh_goals(self):
        """Обновление списка целей"""
        for item in self.goals_tree.get_children():
            self.goals_tree.delete(item)

        goals = database.get_all_goals("iom.db")
        for goal in goals:
            self.goals_tree.insert('', 'end', values=goal)

    def add_goal(self):
        """Открыть окно добавления цели"""
        self._open_goal_window()

    def edit_goal(self):
        """Открыть окно редактирования цели"""
        selected = self.goals_tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите цель для редактирования")
            return

        item = self.goals_tree.item(selected[0])
        goal_id = item['values'][0]
        self._open_goal_window(goal_id)

    def delete_goal(self):
        """Удаление цели"""
        selected = self.goals_tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите цель для удаления")
            return

        item = self.goals_tree.item(selected[0])
        goal_id = item['values'][0]

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            database.delete_goal("iom.db", goal_id)
            self.refresh_goals()
            self.check_all_achievements()
            self.update_stats()
            self.update_semester_progress_auto()

    def _open_goal_window(self, goal_id=None):
        """Общее окно для добавления/редактирования цели"""
        is_edit = goal_id is not None
        goal_data = None

        if is_edit:
            goal_data = database.get_goal_by_id("iom.db", goal_id)
            if not goal_data:
                return

        window = tk.Toplevel(self.root)
        window.title("Редактировать цель" if is_edit else "Добавить цель")
        window.geometry("800x600")

        row = 0

        # Название
        ttk.Label(window, text="Название:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        name_entry = ttk.Entry(window, width=50)
        name_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        if is_edit:
            name_entry.insert(0, goal_data[1])
        row += 1

        # Тип
        ttk.Label(window, text="Тип:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        type_combo = ttk.Combobox(window,
                                   values=['Курс', 'Проект', 'Самообразование', 'Семинар', 'Другое'], width=47)
        type_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        if is_edit:
            type_combo.set(goal_data[2])
        row += 1

        # Статус
        ttk.Label(window, text="Статус:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        status_combo = ttk.Combobox(window,
                                     values=['Новая', 'В процессе', 'Завершена', 'Отменена'], width=47)
        status_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        status_combo.set(goal_data[3] if is_edit else 'Новая')
        row += 1

        # Даты
        ttk.Label(window, text="Плановая дата (ГГГГ-ММ-ДД):").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        plan_date_entry = ttk.Entry(window, width=20)
        plan_date_entry.grid(row=row, column=1, padx=5, pady=5, sticky='w')
        if is_edit and goal_data[4]:
            plan_date_entry.insert(0, goal_data[4])

        ttk.Label(window, text="Фактическая дата (ГГГГ-ММ-ДД):").grid(row=row, column=1, sticky='e', padx=20)
        fact_date_entry = ttk.Entry(window, width=20)
        fact_date_entry.grid(row=row, column=2, padx=5, pady=5, sticky='w')
        if is_edit and goal_data[5]:
            fact_date_entry.insert(0, goal_data[5])
        row += 1

        # Темп
        ttk.Label(window, text="Темп:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        temp_entry = ttk.Entry(window, width=50)
        temp_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        if is_edit and goal_data[6]:
            temp_entry.insert(0, goal_data[6])
        row += 1

        # Навыки
        ttk.Label(window, text="Навыки (до 3, через запятую):").grid(row=row, column=0, sticky='w', padx=5, pady=5)

        all_skills = database.get_all_skills("iom.db")
        skills_entry = ttk.Combobox(window, values=all_skills, width=50)

        if is_edit:
            goal_skills = database.get_goal_skills("iom.db", goal_id)
            skills_entry.set(', '.join(goal_skills))
        skills_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        # Компетенции
        ttk.Label(window, text="Компетенции (выберите 1-3):").grid(row=row, column=0, sticky='nw', padx=5, pady=5)
        comp_frame = ttk.Frame(window)
        comp_frame.grid(row=row, column=1, padx=5, pady=5, columnspan=2, sticky='w')

        all_comps = database.get_all_competencies("iom.db")
        goal_comps = database.get_goal_competencies("iom.db", goal_id) if is_edit else []
        goal_comp_dict = {comp[0]: comp[2] for comp in goal_comps}

        competencies_vars = []
        for comp_id, comp_name in all_comps:
            var = tk.BooleanVar(value=(comp_id in goal_comp_dict))
            cb = ttk.Checkbutton(comp_frame, text=comp_name, variable=var)
            cb.pack(anchor='w')

            level_combo = ttk.Combobox(comp_frame, values=['1', '2', '3', '4', '5'], width=5, state='readonly')
            level_combo.set(str(goal_comp_dict.get(comp_id, '3')))
            level_combo.pack(anchor='w', pady=2)
            level_combo.config(state='readonly' if var.get() else 'disabled')

            var.trace('w', lambda *args, v=var, c=level_combo: utils.toggle_level_combo(v, c))

            competencies_vars.append((var, level_combo, comp_id))

        row += 1

        # Описание
        desc_frame = ttk.LabelFrame(window, text="Описание (простая разметка)")
        desc_frame.grid(row=row, column=0, columnspan=3, padx=5, pady=5, sticky='nsew')

        desc_text = tk.Text(desc_frame, width=40, height=8)
        if is_edit and goal_data[7]:
            desc_text.insert("1.0", goal_data[7])
        desc_text.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        preview_btn = ttk.Button(desc_frame, text="Просмотр",
                                 command=lambda: utils.show_markdown_preview(window, desc_text))
        preview_btn.pack(side='left', padx=5)

        window.grid_rowconfigure(row, weight=1)
        window.grid_columnconfigure(1, weight=1)
        row += 1

        # Кнопка сохранения
        def save():
            self._save_goal(
                goal_id if is_edit else None,
                name_entry.get(),
                type_combo.get(),
                status_combo.get(),
                plan_date_entry.get(),
                fact_date_entry.get(),
                temp_entry.get(),
                desc_text.get("1.0", tk.END).strip(),
                skills_entry.get().strip(),
                competencies_vars
            )
            window.destroy()

        ttk.Button(window, text="Сохранить", command=save).grid(row=row, column=2, sticky='e', padx=5, pady=10)

    def _save_goal(self, goal_id, name, goal_type, status, plan_date, fact_date, temp, description,
                   skills_text, competencies_vars):
        """Сохранение цели (общая логика для добавления и редактирования)"""
        if not name or not goal_type or not status:
            messagebox.showerror("Ошибка", "Заполните обязательные поля: Название, Тип, Статус")
            return

        # Обработка навыков
        skills_list = []
        if skills_text:
            skills_list = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
            if len(skills_list) > 3:
                messagebox.showerror("Ошибка", "Можно указать не более 3 навыков")
                return

        # Обработка компетенций
        selected_competencies = []
        for var, level_combo, comp_id in competencies_vars:
            if var.get():
                level = level_combo.get()
                if not level:
                    messagebox.showerror("Ошибка", "Для выбранных компетенций укажите уровень")
                    return
                selected_competencies.append((comp_id, int(level)))

        if len(selected_competencies) > 3:
            messagebox.showerror("Ошибка", "Можно выбрать не более 3 компетенций")
            return

        if len(selected_competencies) < 1:
            messagebox.showerror("Ошибка", "Выберите хотя бы 1 компетенцию")
            return

        # Сохранение в БД
        if goal_id:  # Редактирование
            database.update_goal("iom.db", goal_id, name, goal_type, status, plan_date, fact_date, temp, description)
        else:  # Добавление
            goal_id = database.add_goal("iom.db", name, goal_type, status, plan_date, fact_date, temp, description)

        if not goal_id:
            messagebox.showerror("Ошибка", "Не удалось сохранить цель")
            return

        # Сохраняем навыки (сначала удаляем старые связи)
        with sqlite3.connect('iom.db') as conn:
            c = conn.cursor()
            c.execute("DELETE FROM цель_навыки WHERE цель_id = ?", (goal_id,))

            for skill_name in skills_list:
                skill_id = database.add_skill("iom.db", skill_name)
                if skill_id:
                    c.execute("INSERT INTO цель_навыки (цель_id, навык_id) VALUES (?, ?)", (goal_id, skill_id))

            # Сохраняем компетенции (сначала удаляем старые связи)
            c.execute("DELETE FROM цель_компетенции WHERE цель_id = ?", (goal_id,))
            for comp_id, level in selected_competencies:
                c.execute('''
                    INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень)
                    VALUES (?, ?, ?)
                ''', (goal_id, comp_id, level))

            conn.commit()

        messagebox.showinfo("Успех", "Цель сохранена")
        self.refresh_goals()
        self.check_all_achievements()
        self.update_stats()
        self.update_semester_progress_auto()

    # ============= ВКЛАДКА "МОЙ ПРОФИЛЬ" =============
    def create_profile_tab(self):
        """Создание вкладки профиля со статистикой"""
        title_frame = ttk.Frame(self.tab_profile)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Мой профиль", font=("Arial", 14, "bold")).pack()

        stats_frame = ttk.LabelFrame(self.tab_profile, text="Статистика")
        stats_frame.pack(fill='both', expand=True, padx=10, pady=10)

        # Навыки
        skills_frame = ttk.LabelFrame(stats_frame, text="Навыки")
        skills_frame.pack(fill='both', expand=True, padx=5, pady=5)

        self.skills_tree = ttk.Treeview(skills_frame, columns=('Навык', 'Количество целей'), show='headings', height=8)
        self.skills_tree.heading('Навык', text='Навык')
        self.skills_tree.heading('Количество целей', text='Количество целей')
        self.skills_tree.pack(fill='both', expand=True, padx=5, pady=5)

        # Статистика по типам целей
        types_frame = ttk.LabelFrame(stats_frame, text="Статистика по типам целей")
        types_frame.pack(fill='both', expand=True, padx=5, pady=5)

        self.types_tree = ttk.Treeview(types_frame, columns=('Тип', 'Завершено', 'Всего'), show='headings', height=5)
        self.types_tree.heading('Тип', text='Тип')
        self.types_tree.heading('Завершено', text='Завершено')
        self.types_tree.heading('Всего', text='Всего')
        self.types_tree.pack(fill='both', expand=True, padx=5, pady=5)

        # Процент целей в срок
        timely_frame = ttk.Frame(stats_frame)
        timely_frame.pack(fill='x', padx=5, pady=5)
        self.timely_label = ttk.Label(timely_frame, text="Процент целей, завершённых в срок: 0%")
        self.timely_label.pack()

        # Кнопка обновления
        ttk.Button(stats_frame, text="Обновить статистику", command=self.update_profile_stats).pack(pady=10)

        self.update_profile_stats()

    def update_profile_stats(self):
        """Обновление статистики в профиле"""
        # Очищаем деревья
        for tree in [self.skills_tree, self.types_tree]:
            for item in tree.get_children():
                tree.delete(item)

        with sqlite3.connect('iom.db') as conn:
            c = conn.cursor()

            # Статистика по навыкам
            c.execute('''
                SELECT н.название, COUNT(цн.цель_id) as количество
                FROM навыка н
                LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
                LEFT JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершена'
                GROUP BY н.id
                HAVING количество > 0
            ''')
            skills = c.fetchall()
            for skill in skills:
                self.skills_tree.insert('', 'end', values=skill)

            # Статистика по типам целей
            c.execute('''
                SELECT тип, 
                       SUM(CASE WHEN статус = 'Завершена' THEN 1 ELSE 0 END) as завершено,
                       COUNT(*) as всего
                FROM цели
                GROUP BY тип
            ''')
            types = c.fetchall()
            for type_stat in types:
                self.types_tree.insert('', 'end', values=type_stat)

            # Процент целей, завершённых в срок
            c.execute("SELECT COUNT(*) FROM цели WHERE статус = 'Завершена' AND факт_дата IS NOT NULL")
            completed_total = c.fetchone()[0]

            c.execute('''
                SELECT COUNT(*) FROM цели 
                WHERE статус = 'Завершена' 
                AND факт_дата IS NOT NULL 
                AND план_дата IS NOT NULL
                AND факт_дата <= план_дата
            ''')
            timely_completed = c.fetchone()[0]

            if completed_total > 0:
                percentage = (timely_completed / completed_total) * 100
                self.timely_label.config(text=f"Процент целей, завершённых в срок: {percentage:.1f}%")
            else:
                self.timely_label.config(text="Процент целей, завершённых в срок: 0%")

    # ============= ВКЛАДКА "КОМПЕТЕНЦИИ" =============
    def create_competencies_tab(self):
        """Создание вкладки компетенций"""
        title_frame = ttk.Frame(self.tab_competencies)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Компетенции", font=("Arial", 14, "bold")).pack()

        comp_frame = ttk.LabelFrame(self.tab_competencies, text="Компетенции и уровни")
        comp_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.competencies_tree = ttk.Treeview(comp_frame, columns=('Компетенция', 'Категория', 'Средний уровень'),
                                              show='headings', height=10)
        self.competencies_tree.heading('Компетенция', text='Компетенция')
        self.competencies_tree.heading('Категория', text='Категория')
        self.competencies_tree.heading('Средний уровень', text='Средний уровень')
        self.competencies_tree.pack(fill='both', expand=True, padx=5, pady=5)

        weak_frame = ttk.LabelFrame(self.tab_competencies, text="Слабые зоны (уровень < 3)")
        weak_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.weak_zones_text = tk.Text(weak_frame, height=5, width=50)
        self.weak_zones_text.pack(fill='both', expand=True, padx=5, pady=5)

        rec_frame = ttk.LabelFrame(self.tab_competencies, text="Рекомендации")
        rec_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.recommendations_text = tk.Text(rec_frame, height=5, width=50)
        self.recommendations_text.pack(fill='both', expand=True, padx=5, pady=5)

        ttk.Button(self.tab_competencies, text="Обновить компетенции", command=self.update_competencies_stats).pack(
            pady=10)

        self.update_competencies_stats()

    def update_competencies_stats(self):
        """Обновление статистики компетенций"""
        for item in self.competencies_tree.get_children():
            self.competencies_tree.delete(item)

        self.weak_zones_text.delete('1.0', tk.END)
        self.recommendations_text.delete('1.0', tk.END)

        comp_stats = database.get_competency_averages("iom.db")

        weak_zones = []
        recommendations = []

        for comp in comp_stats:
            self.competencies_tree.insert('', 'end', values=comp)

            if comp[2] is not None and comp[2] < 3:
                weak_zones.append(f"{comp[0]} - {comp[2]}")

                if comp[0] == "Презентация результатов":
                    recommendations.append(
                        "Вы почти не развиваете компетенцию 'Презентация результатов'. Рекомендуем выступить на студенческой конференции.")
                elif comp[0] == "Работа с БД":
                    recommendations.append(
                        "Для развития компетенции 'Работа с БД' пройдите курс по SQL или поработайте над проектом с базами данных.")
                elif comp[0] == "Управление проектами":
                    recommendations.append(
                        "Для развития 'Управления проектами' возьмите на себя роль тимлида в учебном проекте.")
                else:
                    recommendations.append(
                        f"Для развития компетенции '{comp[0]}' рекомендуется выполнить практические задания.")

        if weak_zones:
            self.weak_zones_text.insert('1.0', '\n'.join(weak_zones))
        else:
            self.weak_zones_text.insert('1.0', "Слабых зон не обнаружено")

        if recommendations:
            self.recommendations_text.insert('1.0', '\n\n'.join(recommendations))
        else:
            self.recommendations_text.insert('1.0',
                                             "Все компетенции развиваются хорошо. Продолжайте в том же духе!")

    # ============= ВКЛАДКА "ДОСТИЖЕНИЯ" =============
    def create_achievements_tab(self):
        """Создание вкладки достижений"""
        title_frame = ttk.Frame(self.tab_achievements)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Достижения", font=("Arial", 14, "bold")).pack()

        self.achievements_tree = ttk.Treeview(self.tab_achievements,
                                              columns=('Получено', 'Название', 'Описание'), show='headings',
                                              height=10)
        self.achievements_tree.heading('Получено', text='Получено')
        self.achievements_tree.heading('Название', text='Название')
        self.achievements_tree.heading('Описание', text='Описание')

        self.achievements_tree.column('Получено', width=80)
        self.achievements_tree.column('Название', width=150)
        self.achievements_tree.column('Описание', width=400)

        self.achievements_tree.pack(fill='both', expand=True, padx=10, pady=10)

        ttk.Button(self.tab_achievements, text="Обновить достижения",
                   command=self.update_achievements_list).pack(pady=10)

        self.update_achievements_list()

    def update_achievements_list(self):
        """Обновление списка достижений"""
        for item in self.achievements_tree.get_children():
            self.achievements_tree.delete(item)

        achievements = database.get_all_achievements("iom.db")
        for ach in achievements:
            status = "✓ Получено" if ach[3] == 1 else "✗ Не получено"
            self.achievements_tree.insert('', 'end', values=(status, ach[1], ach[2]))

    def check_all_achievements(self):
        """Проверка всех достижений"""
        unlocked = database.check_achievements("iom.db")
        self.update_achievements_list()

    # ============= ВКЛАДКА "ЦЕЛИ НА СЕМЕСТР" =============
    def create_semester_tab(self):
        """Создание вкладки целей на семестр"""
        title_frame = ttk.Frame(self.tab_semester)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Цели на семестр", font=("Arial", 14, "bold")).pack()

        goals_frame = ttk.LabelFrame(self.tab_semester, text="Цели семестра")
        goals_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.semester_tree = ttk.Treeview(goals_frame, columns=('ID', 'Цель', 'Тип', 'Прогресс'),
                                          show='headings', height=10)
        self.semester_tree.heading('ID', text='ID')
        self.semester_tree.heading('Цель', text='Цель')
        self.semester_tree.heading('Тип', text='Тип')
        self.semester_tree.heading('Прогресс', text='Прогресс')
        self.semester_tree.pack(fill='both', expand=True, padx=5, pady=5)

        btn_frame = ttk.Frame(goals_frame)
        btn_frame.pack(fill='x', pady=5)

        ttk.Button(btn_frame, text="Добавить цель", command=self.add_semester_goal).pack(side='left',
                                                                                         padx=2)
        ttk.Button(btn_frame, text="Удалить цель", command=self.delete_semester_goal).pack(side='left',
                                                                                           padx=2)
        ttk.Button(btn_frame, text="Обновить", command=self.update_semester_progress_auto).pack(
            side='left', padx=2)

        ttk.Button(self.tab_semester, text="Сформировать отчёт", command=self.export_to_word).pack(pady=10)

        self.update_semester_progress_auto()

    def add_semester_goal(self):
        """Добавление цели на семестр"""
        add_window = tk.Toplevel(self.root)
        add_window.title("Добавить цель на семестр")

        ttk.Label(add_window, text="Текст цели:").grid(row=0, column=0, sticky='w', padx=5, pady=5)
        goal_entry = ttk.Entry(add_window, width=40)
        goal_entry.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(add_window, text="Тип цели:").grid(row=1, column=0, sticky='w', padx=5, pady=5)
        type_combo = ttk.Combobox(add_window, values=['Количество', 'Повышение компетенции', 'Другое'])
        type_combo.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(add_window, text="Параметр (опционально):").grid(row=2, column=0, sticky='w', padx=5,
                                                                   pady=5)
        param_entry = ttk.Entry(add_window, width=40)
        param_entry.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(add_window, text="Целевой прогресс:").grid(row=3, column=0, sticky='w', padx=5, pady=5)
        target_spinbox = ttk.Spinbox(add_window, from_=1, to=100, width=10)
        target_spinbox.grid(row=3, column=1, padx=5, pady=5)
        target_spinbox.set(1)

        def save_semester_goal():
            database.add_semester_goal(
                "iom.db",
                goal_entry.get(),
                type_combo.get(),
                param_entry.get(),
                int(target_spinbox.get())
            )
            messagebox.showinfo("Успех", "Цель на семестр добавлена")
            add_window.destroy()
            self.update_semester_progress_auto()

        ttk.Button(add_window, text="Сохранить", command=save_semester_goal).grid(row=4, column=1,
                                                                                  sticky='e', padx=5,
                                                                                  pady=10)

    def delete_semester_goal(self):
        """Удаление цели на семестр"""
        selected = self.semester_tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите цель для удаления")
            return

        item = self.semester_tree.item(selected[0])
        goal_id = item['values'][0]

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            database.delete_semester_goal("iom.db", goal_id)
            self.update_semester_progress_auto()

    def update_semester_progress_auto(self):
        """Автоматическое обновление прогресса целей на семестр"""
        utils.calculate_semester_progress('iom.db')
        self._refresh_semester_goals()

    def _refresh_semester_goals(self):
        """Обновление отображения целей на семестр"""
        for item in self.semester_tree.get_children():
            self.semester_tree.delete(item)

        goals = database.get_semester_goals("iom.db")
        for goal in goals:
            progress_text = f"{goal[3]} из {goal[4]}"
            self.semester_tree.insert('', 'end',
                                      values=(goal[0], goal[1], goal[2], progress_text))

    # ============= ВКЛАДКА "НАСТРОЙКИ" =============
    def create_settings_tab(self):
        """Создание вкладки настроек"""
        title_frame = ttk.Frame(self.tab_settings)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Настройки", font=("Arial", 14, "bold")).pack()

        spec_frame = ttk.LabelFrame(self.tab_settings, text="Специальность")
        spec_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(spec_frame, text="Выберите специальность:").pack(anchor='w', padx=5, pady=5)

        self.specialty_combo = ttk.Combobox(spec_frame,
                                            values=['Информационные системы',
                                                    'Программная инженерия',
                                                    'Прикладная информатика',
                                                    'Другая'])
        self.specialty_combo.pack(fill='x', padx=5, pady=5)
        self.specialty_combo.set('Информационные системы')

        ttk.Button(spec_frame, text="Сохранить специальность",
                   command=self.save_specialty).pack(pady=5)

        db_frame = ttk.LabelFrame(self.tab_settings, text="База данных")
        db_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(db_frame, text="Текущая БД: SQLite (iom.db)").pack(anchor='w', padx=5, pady=5)

        ttk.Button(db_frame, text="Очистить все данные",
                   command=self.clear_all_data).pack(pady=5)

        info_frame = ttk.LabelFrame(self.tab_settings, text="О программе")
        info_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(info_frame, text="Планировщик индивидуального образовательного маршрута").pack(
            anchor='w', padx=5, pady=2)
        ttk.Label(info_frame, text="Версия 1.0").pack(anchor='w', padx=5, pady=2)
        ttk.Label(info_frame, text="Работает автономно, без подключения к интернету").pack(anchor='w',
                                                                                           padx=5,
                                                                                           pady=2)

    def save_specialty(self):
        """Сохранение специальности"""
        specialty = self.specialty_combo.get()
        messagebox.showinfo("Сохранено", f"Специальность '{specialty}' сохранена")

    def clear_all_data(self):
        """Очистка всех данных (для отладки)"""
        if messagebox.askyesno("Подтверждение",
                               "Вы уверены, что хотите удалить все данные?\nЭто действие нельзя отменить."):
            with sqlite3.connect('iom.db') as conn:
                c = conn.cursor()

                tables = ['цели', 'навыка', 'цель_навыки', 'компетенции',
                          'цель_компетенции', 'цель_каса']

                for table in tables:
                    c.execute(f"DELETE FROM {table}")

                c.execute("UPDATE достижения SET получено = 0")
                conn.commit()

            # Перезагружаем компетенции из JSON
            database.load_competencies_to_db()

            messagebox.showinfo("Очищено", "Все данные удалены")

            self.refresh_goals()
            self.update_profile_stats()
            self.update_competencies_stats()
            self.update_achievements_list()
            self.update_semester_progress_auto()

    # ============= ОБЩИЕ МЕТОДЫ =============
    def update_stats(self):
        """Обновление всей статистики"""
        self.update_profile_stats()
        self.update_competencies_stats()

    def export_to_word(self):
        """Экспорт отчёта в Word"""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            title = doc.add_heading('Индивидуальный образовательный маршрут', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Отчёт сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph()

            doc.add_heading('Цели', level=1)

            with sqlite3.connect('iom.db') as conn:
                c = conn.cursor()
                c.execute("SELECT * FROM цели ORDER BY статус, план_дата")
                goals = c.fetchall()

                if goals:
                    for goal in goals:
                        doc.add_heading(goal[1], level=2)
                        doc.add_paragraph(f"Тип: {goal[2]}")
                        doc.add_paragraph(f"Статус: {goal[3]}")
                        doc.add_paragraph(f"Плановая дата: {goal[4] or 'Не указана'}")
                        doc.add_paragraph(f"Фактическая дата: {goal[5] or 'Не указана'}")
                        if goal[6]:
                            doc.add_paragraph(f"Темп: {goal[6]}")

                        if goal[7]:
                            utils.add_formatted_text_to_doc(doc, goal[7])

                        doc.add_paragraph()
                else:
                    doc.add_paragraph("Цели не добавлены")

                doc.add_heading('Навыки', level=1)
                c.execute('''
                    SELECT н.название, COUNT(цн.цель_id) as количество
                    FROM навыка н
                    LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
                    LEFT JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершена'
                    GROUP BY н.id
                    HAVING количество > 0
                ''')
                skills = c.fetchall()

                if skills:
                    for skill in skills:
                        doc.add_paragraph(f"{skill[0]} — {skill[1]} цели", style='List Bullet')
                else:
                    doc.add_paragraph("Навыки не указаны")

                doc.add_heading('Компетенции', level=1)

                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Компетенция'
                hdr_cells[1].text = 'Категория'
                hdr_cells[2].text = 'Средний уровень'

                comp_stats = database.get_competency_averages("iom.db")

                for comp in comp_stats:
                    row_cells = table.add_row().cells
                    row_cells[0].text = comp[0] or ''
                    row_cells[1].text = comp[1] or ''
                    row_cells[2].text = str(comp[2]) if comp[2] else 'Нет данных'

                doc.add_heading('Слабые зоны', level=1)
                weak_zones = [comp for comp in comp_stats if comp[2] is not None and comp[2] < 3]

                if weak_zones:
                    for zone in weak_zones:
                        doc.add_paragraph(f"{zone[0]} — уровень {zone[2]}", style='List Bullet')
                else:
                    doc.add_paragraph("Слабых зон не обнаружено")

                doc.add_heading('Рекомендации', level=1)

                recommendations = []
                for comp in comp_stats:
                    if comp[2] is not None and comp[2] < 3:
                        if comp[0] == "Презентация результатов":
                            recommendations.append(
                                "Вы почти не развиваете компетенцию 'Презентация результатов'. Рекомендуем выступить на студенческой конференции.")
                        elif comp[0] == "Работа с БД":
                            recommendations.append(
                                "Для развития компетенции 'Работа с БД' пройдите курс по SQL или поработайте над проектом с базами данных.")
                        elif comp[0] == "Управление проектами":
                            recommendations.append(
                                "Для развития 'Управления проектами' возьмите на себя роль тимлида в учебном проекте.")

                if recommendations:
                    for rec in recommendations:
                        doc.add_paragraph(rec, style='List Bullet')
                else:
                    doc.add_paragraph("Все компетенции развиваются хорошо. Продолжайте в том же духе!")

                doc.add_heading('Достижения', level=1)
                achievements = database.get_all_achievements("iom.db")
                obtained = [a for a in achievements if a[3] == 1]

                if obtained:
                    for ach in obtained:
                        doc.add_paragraph(f"{ach[1]} — {ach[2]}", style='List Bullet')
                else:
                    doc.add_paragraph("Достижения не получены")

                doc.add_heading('Цели на семестр', level=1)
                semester_goals = database.get_semester_goals("iom.db")

                if semester_goals:
                    for goal in semester_goals:
                        doc.add_paragraph(f"{goal[1]} — {goal[3]} из {goal[4]}", style='List Bullet')
                else:
                    doc.add_paragraph("Цели на семестр не установлены")

            filename = f"Отчет_ИОМ_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            messagebox.showinfo("Успех", f"Отчёт сохранён в файл: {filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать отчёт: {str(e)}")