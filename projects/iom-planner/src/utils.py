"""
Вспомогательные функции (разметка, экспорт, превью)
"""
import tkinter as tk
from tkinter import ttk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from datetime import datetime
import sqlite3
from typing import Optional


def parse_simple_markdown(text: str) -> str:
    """Преобразование простой разметки в текст для отображения в GUI"""
    lines = text.split('\n')
    result_lines = []
    for line in lines:
        if line.startswith('- '):
            result_lines.append('• ' + line[2:])
        elif line.startswith('# '):
            result_lines.append('ЗАГОЛОВОК: ' + line[2:])
        else:
            # Обработка **текст** (жирный)
            line = line.replace('**', '')
            # Обработка [текст](ссылка)
            if '[' in line and ']' in line and '(' in line and ')' in line:
                start = line.find('[')
                end = line.find(']')
                link_start = line.find('(')
                link_end = line.find(')')
                text = line[start + 1:end]
                link = line[link_start + 1:link_end]
                line = line[:start] + text + ' (' + link + ')' + line[link_end + 1:]
            result_lines.append(line)
    return '\n'.join(result_lines)


def show_markdown_preview(parent, text_widget):
    """Показ окна предпросмотра разметки"""
    preview_window = tk.Toplevel(parent)
    preview_window.title("Предпросмотр разметки")
    preview_window.geometry("600x400")

    preview_text = tk.Text(preview_window, wrap='word', font=('Arial', 10))
    preview_text.pack(fill='both', expand=True, padx=10, pady=10)

    text = text_widget.get("1.0", tk.END)
    lines = text.split('\n')

    for line in lines:
        line = line.rstrip()
        if line.startswith('- '):
            preview_text.insert(tk.END, '• ' + line[2:] + '\n')
        elif line.startswith('# '):
            preview_text.insert(tk.END, 'ЗАГОЛОВОК: ' + line[2:] + '\n', 'header')
        elif '**' in line:
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    preview_text.insert(tk.END, part.upper(), 'bold')
                else:
                    preview_text.insert(tk.END, part)
            preview_text.insert(tk.END, '\n')
        elif '[' in line and ']' in line and '(' in line and ')' in line:
            start = line.find('[')
            end = line.find(']')
            link_start = line.find('(')
            link_end = line.find(')')
            link_text = line[start + 1:end]
            link_url = line[link_start + 1:link_end]
            preview_text.insert(tk.END, line[:start])
            preview_text.insert(tk.END, link_text + ' ', 'link')
            preview_text.insert(tk.END, f'({link_url})')
            preview_text.insert(tk.END, line[link_end + 1:] + '\n')
        else:
            preview_text.insert(tk.END, line + '\n')

    preview_text.tag_configure('header', font=('Arial', 12, 'bold'))
    preview_text.tag_configure('bold', font=('Arial', 10, 'bold'))
    preview_text.tag_configure('link', foreground='blue', underline=True)

    preview_text.config(state='disabled')


def toggle_level_combo(var, combo):
    """Включение/отключение комбобокса уровня компетенции"""
    combo.config(state='readonly' if var.get() else 'disabled')


def add_formatted_text_to_doc(doc, text):
    """Добавление форматированного текста в Word документ"""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], level=2)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:])
        elif '**' in line:
            parts = line.split('**')
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
        elif '*' in line:
            parts = line.split('*')
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.italic = True
                else:
                    p.add_run(part)
        elif '__' in line:
            parts = line.split('__')
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.italic = True
                else:
                    p.add_run(part)
        else:
            if '[' in line and ']' in line and '(' in line and ')' in line:
                start = line.find('[')
                end = line.find(']')
                link_start = line.find('(')
                link_end = line.find(')')
                link_text = line[start + 1:end]
                link_url = line[link_start + 1:link_end]

                p = doc.add_paragraph()
                p.add_run(line[:start])

                # Создаем гиперссылку
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), f'rId{doc.part.next_id}')

                run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')

                # Синий цвет
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '0000FF')
                rPr.append(color)

                # Подчеркивание
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)

                run.append(rPr)
                text_element = OxmlElement('w:t')
                text_element.text = link_text
                run.append(text_element)

                hyperlink.append(run)
                p._element.append(hyperlink)

                # Добавляем остальной текст
                if line[link_end + 1:]:
                    p.add_run(line[link_end + 1:])
            else:
                doc.add_paragraph(line)


def calculate_semester_progress(db_path: str) -> None:
    """Автоматический расчет прогресса целей на семестр"""
    try:
        with sqlite3.connect(db_path) as conn:
            c = conn.cursor()

            c.execute("SELECT id, текст_цели, тип_цели, параметр, целевой_прогресс FROM цель_каса")
            semester_goals = c.fetchall()

            for goal in semester_goals:
                goal_id, text, goal_type, param, target = goal
                current = 0

                if goal_type == 'Количество':
                    if 'курс' in text.lower() or (param and 'курс' in param.lower()):
                        c.execute("SELECT COUNT(*) FROM цели WHERE тип = 'Курс' AND статус = 'Завершена'")
                        current = c.fetchone()[0]
                    elif 'проект' in text.lower() or (param and 'проект' in param.lower()):
                        c.execute("SELECT COUNT(*) FROM цели WHERE тип = 'Проект' AND статус = 'Завершена'")
                        current = c.fetchone()[0]
                    elif 'семинар' in text.lower() or (param and 'семинар' in param.lower()):
                        c.execute("SELECT COUNT(*) FROM цели WHERE тип = 'Семинар' AND статус = 'Завершена'")
                        current = c.fetchone()[0]
                    else:
                        c.execute("SELECT COUNT(*) FROM цели WHERE статус = 'Завершена'")
                        current = c.fetchone()[0]

                elif goal_type == 'Повышение компетенции' and param:
                    c.execute('''
                        SELECT ROUND(AVG(цк.уровень), 0)
                        FROM компетенции к
                        JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                        JOIN цели ц ON цк.цель_id = ц.id AND ц.статус = 'Завершена'
                        WHERE к.название = ?
                    ''', (param,))
                    result = c.fetchone()
                    current = int(result[0]) if result and result[0] else 0

                c.execute('''
                    UPDATE цель_каса 
                    SET текущий_прогресс = ?
                    WHERE id = ?
                ''', (current, goal_id))

            conn.commit()
    except Exception as e:
        print(f"❌ Ошибка расчета прогресса: {e}")