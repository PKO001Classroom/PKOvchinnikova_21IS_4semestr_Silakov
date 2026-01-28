import sqlite3
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
import json
import os


# ===== 1. ИНИЦИАЛИЗАЦИЯ БАЗЫ ДАННЫХ =====
def init_db():
    """Создание базы данных и таблицы, если их нет"""
    conn = sqlite3.connect("достижения.db")
    cur = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS достижения (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            название TEXT NOT NULL,
            дата TEXT NOT NULL,
            тип TEXT,
            уровень TEXT,
            описание TEXT
        )
    ''')
    conn.commit()
    conn.close()
    print("База данных инициализирована")


# ===== 2. ЗАГРУЗКА ТИПОВ ИЗ JSON-ФАЙЛА =====
def load_types():
    """Загрузка типов достижений из файла types.json"""
    try:
        if os.path.exists("types.json"):
            with open("types.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list) and data:
                    return data
        # Если файла нет или он пустой - возвращаем значения по умолчанию
        return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]
    except (json.JSONDecodeError, Exception) as e:
        print(f"Ошибка загрузки types.json: {e}")
        return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]


# ===== 3. СОХРАНЕНИЕ ДАННЫХ В БАЗУ =====
def save_to_db(name, date, typ, level, desc):
    """Сохранение достижения в базу данных"""
    try:
        conn = sqlite3.connect("достижения.db")
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO достижения (название, дата, тип, уровень, описание) VALUES (?, ?, ?, ?, ?)",
            (name, date, typ, level, desc)
        )
        conn.commit()
        conn.close()
        print(f"Успешно сохранено: {name}")
        return True
    except Exception as e:
        print(f"Ошибка сохранения в БД: {e}")
        return False


# ===== 4. ЗАГРУЗКА ЗАПИСЕЙ ИЗ БАЗЫ =====
def load_records():
    """Загрузка всех записей из базы данных"""
    try:
        conn = sqlite3.connect("достижения.db")
        cur = conn.cursor()
        cur.execute("SELECT id, название, дата, тип, уровень, описание FROM достижения ORDER BY дата DESC")
        rows = cur.fetchall()
        conn.close()
        return rows
    except Exception as e:
        print(f"Ошибка загрузки из БД: {e}")
        return []


# ===== 5. ЭКСПОРТ В WORD =====
def export_to_word():
    """Экспорт всех достижений в Word-документ"""
    try:
        doc = Document()
        doc.add_heading("Личные учебные достижения", 0)

        records = load_records()

        if not records:
            doc.add_paragraph("Нет сохранённых достижений.")
        else:
            for record in records:
                id_num, name, date, typ, level, desc = record

                # Добавляем достижение
                p = doc.add_paragraph()

                # Название - жирным
                title_run = p.add_run(f"{name}")
                title_run.bold = True

                # Дата - курсивом
                date_run = p.add_run(f" — {date}")
                date_run.italic = True

                # Тип и уровень
                p.add_run(f" ({typ}, {level})")

                # Описание (если есть)
                if desc and desc.strip():
                    doc.add_paragraph(f"Описание: {desc}")

                # Разделительная линия
                doc.add_paragraph()

        # Сохраняем документ
        filename = "достижения.docx"
        doc.save(filename)
        print(f"Документ сохранён: {filename}")

        # Показываем сообщение об успехе
        messagebox.showinfo("Успех", f"Документ '{filename}' успешно сохранён в папке с программой!")

    except Exception as e:
        messagebox.showerror("Ошибка", f"Не удалось сохранить документ:\n{str(e)}")


# ===== 6. СОЗДАНИЕ ГРАФИЧЕСКОГО ИНТЕРФЕЙСА =====
class AchievementsApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Учёт личных достижений")
        self.root.geometry("900x700")

        # Инициализация базы данных
        init_db()

        # Загрузка типов
        self.available_types = load_types()

        # Переменная для хранения записей
        self.current_records = []

        # Создаем Notebook (вкладки)
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Создаем вкладки
        self.tab_add = tk.Frame(self.notebook)
        self.tab_list = tk.Frame(self.notebook)

        self.notebook.add(self.tab_add, text="➕ Добавить достижение")
        self.notebook.add(self.tab_list, text="📋 Мои достижения")

        # Инициализируем вкладки
        self.create_add_tab()
        self.create_list_tab()

        # Статус бар
        self.status_bar = tk.Label(root, text="Готово", bd=1, relief=tk.SUNKEN, anchor=tk.W, padx=10)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        # Обновляем список при запуске
        self.refresh_list()

    def create_add_tab(self):
        """Создание вкладки для добавления достижений"""
        # Заголовок
        tk.Label(self.tab_add, text="Добавление нового достижения",
                 font=("Arial", 14, "bold"), fg="#2196F3").pack(pady=(10, 20))

        # Фрейм для формы
        form_frame = tk.Frame(self.tab_add)
        form_frame.pack(padx=20, pady=10)

        # Название
        tk.Label(form_frame, text="Название достижения:",
                 font=("Arial", 10, "bold")).grid(row=0, column=0, sticky="w", pady=(0, 10))
        self.name_entry = tk.Entry(form_frame, width=60, font=("Arial", 10))
        self.name_entry.grid(row=0, column=1, padx=(10, 0), pady=(0, 10))

        # Дата
        tk.Label(form_frame, text="Дата (ГГГГ-ММ-ДД):",
                 font=("Arial", 10, "bold")).grid(row=1, column=0, sticky="w", pady=(0, 10))
        self.date_entry = tk.Entry(form_frame, width=60, font=("Arial", 10))
        self.date_entry.grid(row=1, column=1, padx=(10, 0), pady=(0, 10))
        self.date_entry.insert(0, "2024-01-01")

        # Тип
        tk.Label(form_frame, text="Тип:",
                 font=("Arial", 10, "bold")).grid(row=2, column=0, sticky="w", pady=(0, 10))
        self.type_combo = ttk.Combobox(form_frame, values=self.available_types,
                                       state="readonly", width=58, font=("Arial", 10))
        self.type_combo.grid(row=2, column=1, padx=(10, 0), pady=(0, 10))
        if self.available_types:
            self.type_combo.set(self.available_types[0])

        # Уровень
        tk.Label(form_frame, text="Уровень:",
                 font=("Arial", 10, "bold")).grid(row=3, column=0, sticky="w", pady=(0, 10))
        self.level_combo = ttk.Combobox(form_frame,
                                        values=["Школьный", "Городской", "Региональный", "Всероссийский",
                                                "Международный", "Другой"],
                                        state="readonly", width=58, font=("Arial", 10))
        self.level_combo.grid(row=3, column=1, padx=(10, 0), pady=(0, 10))
        self.level_combo.set("Школьный")

        # Описание
        tk.Label(form_frame, text="Описание:",
                 font=("Arial", 10, "bold")).grid(row=4, column=0, sticky="nw", pady=(0, 10))

        desc_frame = tk.Frame(form_frame)
        desc_frame.grid(row=4, column=1, padx=(10, 0), pady=(0, 10), sticky="nsew")

        self.desc_text = tk.Text(desc_frame, width=45, height=10, font=("Arial", 10), wrap=tk.WORD)
        self.desc_text.pack(side=tk.LEFT, fill='both', expand=True)

        desc_scrollbar = tk.Scrollbar(desc_frame, command=self.desc_text.yview)
        desc_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.desc_text.config(yscrollcommand=desc_scrollbar.set)

        # Кнопка сохранения
        btn_frame = tk.Frame(self.tab_add)
        btn_frame.pack(pady=20)

        self.save_btn = tk.Button(btn_frame, text="💾 Сохранить достижение",
                                  command=self.on_save, bg="#4CAF50", fg="white",
                                  font=("Arial", 11, "bold"), padx=30, pady=10)
        self.save_btn.pack()

    def create_list_tab(self):
        """Создание вкладки для просмотра достижений"""
        # Заголовок
        tk.Label(self.tab_list, text="Список ваших достижений",
                 font=("Arial", 14, "bold"), fg="#2196F3").pack(pady=(10, 5))

        # Фрейм для списка
        list_frame = tk.Frame(self.tab_list)
        list_frame.pack(pady=10, padx=10, fill='both', expand=True)

        # Listbox с прокруткой
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.listbox = tk.Listbox(list_frame, width=100, height=25,
                                  yscrollcommand=scrollbar.set,
                                  font=("Consolas", 10),
                                  selectbackground="#2196F3",
                                  selectforeground="white")
        self.listbox.pack(side=tk.LEFT, fill='both', expand=True)

        scrollbar.config(command=self.listbox.yview)

        # Фрейм для кнопок
        btn_frame = tk.Frame(self.tab_list)
        btn_frame.pack(pady=10)

        # Кнопка обновления
        tk.Button(btn_frame, text="🔄 Обновить список",
                  command=self.refresh_list,
                  bg="#FF9800", fg="white",
                  font=("Arial", 10), padx=20, pady=5).pack(side=tk.LEFT, padx=5)

        # Кнопка просмотра деталей
        tk.Button(btn_frame, text="👁️ Просмотреть детали",
                  command=self.show_details,
                  bg="#2196F3", fg="white",
                  font=("Arial", 10), padx=20, pady=5).pack(side=tk.LEFT, padx=5)

        # Кнопка экспорта
        tk.Button(btn_frame, text="📄 Экспорт в Word",
                  command=export_to_word,
                  bg="#4CAF50", fg="white",
                  font=("Arial", 10), padx=20, pady=5).pack(side=tk.LEFT, padx=5)

        # Кнопка удаления
        tk.Button(btn_frame, text="🗑️ Удалить выбранное",
                  command=self.delete_selected,
                  bg="#f44336", fg="white",
                  font=("Arial", 10), padx=20, pady=5).pack(side=tk.LEFT, padx=5)

    def on_save(self):
        """Обработка сохранения достижения"""
        name = self.name_entry.get().strip()
        date = self.date_entry.get().strip()
        typ = self.type_combo.get()
        level = self.level_combo.get()
        desc = self.desc_text.get("1.0", tk.END).strip()

        # Проверка ввода
        if not name:
            messagebox.showwarning("Внимание", "Введите название достижения!")
            self.name_entry.focus()
            return

        if not date:
            messagebox.showwarning("Внимание", "Введите дату!")
            self.date_entry.focus()
            return

        # Сохранение
        if save_to_db(name, date, typ, level, desc):
            messagebox.showinfo("Успех", "Достижение успешно сохранено!")

            # Очистка полей
            self.name_entry.delete(0, tk.END)
            self.date_entry.delete(0, tk.END)
            self.date_entry.insert(0, "2024-01-01")
            if self.available_types:
                self.type_combo.set(self.available_types[0])
            self.level_combo.set("Школьный")
            self.desc_text.delete("1.0", tk.END)

            # Обновление списка
            self.refresh_list()
            self.update_status(f"Достижение сохранено. Всего: {len(self.current_records)}")
        else:
            messagebox.showerror("Ошибка", "Не удалось сохранить достижение!")

    def refresh_list(self):
        """Обновление списка достижений"""
        self.listbox.delete(0, tk.END)
        self.current_records = load_records()

        if not self.current_records:
            self.listbox.insert(tk.END, "Нет сохранённых достижений. Добавьте их на вкладке 'Добавить достижение'.")
            self.listbox.itemconfig(0, fg="gray")
        else:
            for record in self.current_records:
                id_num, name, date, typ, level, desc = record
                # Форматируем строку для отображения
                display_text = f"{date} | {name[:50]}{'...' if len(name) > 50 else ''} | {typ} | {level}"
                self.listbox.insert(tk.END, display_text)

        self.update_status(f"Загружено достижений: {len(self.current_records)}")

    def show_details(self):
        """Показ деталей выбранного достижения"""
        selection = self.listbox.curselection()
        if not selection:
            messagebox.showinfo("Информация", "Выберите запись из списка")
            return

        index = selection[0]
        if index < len(self.current_records):
            id_num, name, date, typ, level, desc = self.current_records[index]

            # Создаем окно с деталями
            details_win = tk.Toplevel(self.root)
            details_win.title(f"Детали: {name}")
            details_win.geometry("700x550")
            details_win.transient(self.root)
            details_win.grab_set()

            # Заголовок
            tk.Label(details_win, text=name, font=("Arial", 16, "bold"),
                     fg="#2196F3", wraplength=650).pack(pady=(20, 10))

            # Фрейм для информации
            info_frame = tk.Frame(details_win)
            info_frame.pack(fill='x', padx=30, pady=10)

            # Дата
            tk.Label(info_frame, text="📅 Дата:", font=("Arial", 11, "bold"),
                     width=10, anchor="w").grid(row=0, column=0, sticky="w", pady=8)
            tk.Label(info_frame, text=date, font=("Arial", 11)).grid(row=0, column=1,
                                                                     sticky="w", pady=8, padx=(10, 0))

            # Тип
            tk.Label(info_frame, text="🏷️ Тип:", font=("Arial", 11, "bold"),
                     width=10, anchor="w").grid(row=1, column=0, sticky="w", pady=8)
            tk.Label(info_frame, text=typ, font=("Arial", 11)).grid(row=1, column=1,
                                                                    sticky="w", pady=8, padx=(10, 0))

            # Уровень
            tk.Label(info_frame, text="📊 Уровень:", font=("Arial", 11, "bold"),
                     width=10, anchor="w").grid(row=2, column=0, sticky="w", pady=8)
            tk.Label(info_frame, text=level, font=("Arial", 11)).grid(row=2, column=1,
                                                                      sticky="w", pady=8, padx=(10, 0))

            # Описание
            tk.Label(details_win, text="📝 Описание:", font=("Arial", 12, "bold")).pack(anchor="w",
                                                                                       padx=30, pady=(20, 5))

            desc_frame = tk.Frame(details_win)
            desc_frame.pack(fill='both', expand=True, padx=30, pady=(0, 20))

            desc_scrollbar = tk.Scrollbar(desc_frame)
            desc_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

            desc_text = tk.Text(desc_frame, wrap=tk.WORD, yscrollcommand=desc_scrollbar.set,
                                font=("Arial", 10), height=15)
            desc_text.pack(side=tk.LEFT, fill='both', expand=True)
            desc_text.insert('1.0', desc if desc and desc.strip() else "Описание отсутствует")
            desc_text.config(state='disabled')

            desc_scrollbar.config(command=desc_text.yview)

            # Кнопка закрытия
            tk.Button(details_win, text="Закрыть", command=details_win.destroy,
                      bg="#f44336", fg="white", font=("Arial", 10), padx=30).pack(pady=10)

    def delete_selected(self):
        """Удаление выбранного достижения"""
        selection = self.listbox.curselection()
        if not selection:
            messagebox.showinfo("Информация", "Выберите запись для удаления")
            return

        index = selection[0]
        if index < len(self.current_records):
            id_num, name, date, typ, level, desc = self.current_records[index]

            # Подтверждение удаления
            confirm = messagebox.askyesno("Подтверждение",
                                          f"Вы уверены, что хотите удалить достижение:\n\n{name}?")

            if confirm:
                try:
                    conn = sqlite3.connect("достижения.db")
                    cur = conn.cursor()
                    cur.execute("DELETE FROM достижения WHERE id = ?", (id_num,))
                    conn.commit()
                    conn.close()

                    messagebox.showinfo("Успех", "Достижение удалено!")
                    self.refresh_list()

                except Exception as e:
                    messagebox.showerror("Ошибка", f"Не удалось удалить достижение:\n{str(e)}")

    def update_status(self, text):
        """Обновление текста в статусной строке"""
        self.status_bar.config(text=text)


# ===== ЗАПУСК ПРОГРАММЫ =====
if __name__ == "__main__":
    try:
        root = tk.Tk()
        app = AchievementsApp(root)
        root.mainloop()
    except Exception as e:
        print(f"Критическая ошибка: {e}")
        messagebox.showerror("Ошибка", f"Не удалось запустить программу:\n{str(e)}")