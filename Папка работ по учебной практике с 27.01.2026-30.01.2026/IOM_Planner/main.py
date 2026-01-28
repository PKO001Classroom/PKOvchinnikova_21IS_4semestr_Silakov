import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3
import json
import os
from datetime import datetime


def init_database():
    conn = sqlite3.connect('iom.db')
    c = conn.cursor()

    c.execute('''
        CREATE TABLE IF NOT EXISTS цели (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            название TEXT NOT NULL,
            тип TEXT NOT NULL,
            статус TEXT DEFAULT 'Новая',
            план_дата TEXT,
            факт_дата TEXT,
            темп TEXT,
            описание TEXT
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS навыка (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            название TEXT UNIQUE NOT NULL
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS цель_навыки (
            цель_id INTEGER,
            навык_id INTEGER,
            FOREIGN KEY (цель_id) REFERENCES цели (id),
            FOREIGN KEY (навык_id) REFERENCES навыка (id)
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS компетенции (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            название TEXT NOT NULL,
            категория TEXT
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS цель_компетенции (
            цель_id INTEGER,
            компетенция_id INTEGER,
            уровень INTEGER CHECK (уровень BETWEEN 1 AND 5),
            FOREIGN KEY (цель_id) REFERENCES цели (id),
            FOREIGN KEY (компетенция_id) REFERENCES компетенции (id)
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS достижения (
            код TEXT PRIMARY KEY,
            название TEXT NOT NULL,
            описание TEXT,
            получено INTEGER DEFAULT 0
        )
    ''')

    c.execute('''
        CREATE TABLE IF NOT EXISTS цель_каса (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            текст_цели TEXT NOT NULL,
            тип_цели TEXT,
            параметр TEXT,
            текущий_прогресс INTEGER DEFAULT 0,
            целевой_прогресс INTEGER NOT NULL
        )
    ''')

    achievements = [
        ('ach1', 'Старт', 'Создана первая цель', 0),
        ('ach2', 'Пунктуальный', 'Три или более завершённых целей в срок', 0),
        ('ach3', 'Многогранный', 'Есть цели минимум трёх разных типов', 0),
        ('ach4', 'Навыковый рост', 'У одного навыка четыре или более связанных завершённых целей', 0),
        ('ach5', 'Планирование', 'Одновременно в статусе "В процессе" пять или более целей', 0)
    ]

    c.execute("SELECT COUNT(*) FROM достижения")
    if c.fetchone()[0] == 0:
        c.executemany("INSERT INTO достижения VALUES (?, ?, ?, ?)", achievements)

    conn.commit()
    conn.close()


def load_competencies_to_db():
    conn = sqlite3.connect('iom.db')
    c = conn.cursor()

    c.execute("SELECT COUNT(*) FROM компетенции")
    if c.fetchone()[0] == 0:
        if os.path.exists('competencies.json'):
            with open('competencies.json', 'r', encoding='utf-8') as f:
                competencies = json.load(f)
                for comp in competencies:
                    c.execute("INSERT INTO компетенции (название, категория) VALUES (?, ?)",
                              (comp['название'], comp['категория']))
        else:
            messagebox.showwarning("Предупреждение", "Файл competencies.json не найден")

    conn.commit()
    conn.close()


class IOMApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Планировщик индивидуального образовательного маршрута")
        self.root.geometry("1200x800")

        init_database()
        load_competencies_to_db()

        self.create_widgets()
        self.check_all_achievements()
        self.update_stats()
        self.update_semester_progress_auto()

    def create_widgets(self):
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill='both', expand=True)

        self.tab_goals = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_goals, text='Мои цели')
        self.create_goals_tab()

        self.tab_profile = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_profile, text='Мой профиль')
        self.create_profile_tab()

        self.tab_competencies = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_competencies, text='Компетенции')
        self.create_competencies_tab()

        self.tab_achievements = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_achievements, text='Достижения')
        self.create_achievements_tab()

        self.tab_semester = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_semester, text='Цели на семестр')
        self.create_semester_tab()

        self.tab_settings = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_settings, text='Настройки')
        self.create_settings_tab()

    def create_goals_tab(self):
        list_frame = ttk.Frame(self.tab_goals)
        list_frame.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        self.goals_tree = ttk.Treeview(list_frame, columns=('ID', 'Название', 'Тип', 'Статус'), show='headings')
        self.goals_tree.heading('ID', text='ID')
        self.goals_tree.heading('Название', text='Название')
        self.goals_tree.heading('Тип', text='Тип')
        self.goals_tree.heading('Статус', text='Статус')
        self.goals_tree.pack(fill='both', expand=True)

        btn_frame = ttk.Frame(list_frame)
        btn_frame.pack(fill='x', pady=5)

        ttk.Button(btn_frame, text='Добавить', command=self.add_goal).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='Редактировать', command=self.edit_goal).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='Удалить', command=self.delete_goal).pack(side='left', padx=2)
        ttk.Button(btn_frame, text='Обновить', command=self.refresh_goals).pack(side='left', padx=2)

        self.refresh_goals()

    def add_goal(self):
        self.add_window = tk.Toplevel(self.root)
        self.add_window.title("Добавить цель")
        self.add_window.geometry("800x600")

        row = 0

        ttk.Label(self.add_window, text="Название:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        self.name_entry = ttk.Entry(self.add_window, width=50)
        self.name_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(self.add_window, text="Тип:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        self.type_combo = ttk.Combobox(self.add_window,
                                       values=['Курс', 'Проект', 'Самообразование', 'Семинар', 'Другое'], width=47)
        self.type_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(self.add_window, text="Статус:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        self.status_combo = ttk.Combobox(self.add_window,
                                         values=['Новая', 'В процессе', 'Завершена', 'Отменена'], width=47)
        self.status_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        self.status_combo.set('Новая')
        row += 1

        ttk.Label(self.add_window, text="Плановая дата (ГГГГ-ММ-ДД):").grid(row=row, column=0, sticky='w', padx=5,
                                                                            pady=5)
        self.plan_date_entry = ttk.Entry(self.add_window, width=20)
        self.plan_date_entry.grid(row=row, column=1, padx=5, pady=5, sticky='w')

        ttk.Label(self.add_window, text="Фактическая дата (ГГГГ-ММ-ДД):").grid(row=row, column=1, sticky='e', padx=20)
        self.fact_date_entry = ttk.Entry(self.add_window, width=20)
        self.fact_date_entry.grid(row=row, column=2, padx=5, pady=5, sticky='w')
        row += 1

        ttk.Label(self.add_window, text="Темп:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        self.temp_entry = ttk.Entry(self.add_window, width=50)
        self.temp_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(self.add_window, text="Навыки (до 3, через запятую):").grid(row=row, column=0, sticky='w', padx=5,
                                                                              pady=5)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT название FROM навыка")
        existing_skills = [skill[0] for skill in c.fetchall()]
        conn.close()

        self.skills_combo = ttk.Combobox(self.add_window, values=existing_skills, width=50)
        self.skills_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(self.add_window, text="Компетенции (выберите 1-3):").grid(row=row, column=0, sticky='nw', padx=5,
                                                                            pady=5)
        comp_frame = ttk.Frame(self.add_window)
        comp_frame.grid(row=row, column=1, padx=5, pady=5, columnspan=2, sticky='w')

        self.competencies_vars = []
        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT id, название FROM компетенции")
        comps = c.fetchall()

        for comp in comps:
            var = tk.BooleanVar()
            cb = ttk.Checkbutton(comp_frame, text=comp[1], variable=var)
            cb.pack(anchor='w')

            level_combo = ttk.Combobox(comp_frame, values=['1', '2', '3', '4', '5'], width=5, state='readonly')
            level_combo.pack(anchor='w', pady=2)
            level_combo.set('3')
            level_combo.config(state='disabled')

            var.trace('w', lambda *args, v=var, c=level_combo: self.toggle_level_combo(v, c))

            self.competencies_vars.append((var, level_combo, comp[0]))

        conn.close()
        row += 1

        desc_frame = ttk.LabelFrame(self.add_window, text="Описание (простая разметка)")
        desc_frame.grid(row=row, column=0, columnspan=3, padx=5, pady=5, sticky='nsew')

        self.desc_text = tk.Text(desc_frame, width=40, height=8)
        self.desc_text.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        preview_btn = ttk.Button(desc_frame, text="Просмотр",
                                 command=lambda: self.show_markdown_preview(self.desc_text))
        preview_btn.pack(side='left', padx=5)

        self.add_window.grid_rowconfigure(row, weight=1)
        self.add_window.grid_columnconfigure(1, weight=1)
        row += 1

        ttk.Button(self.add_window, text="Сохранить", command=self.save_goal).grid(row=row, column=2, sticky='e',
                                                                                   padx=5, pady=10)

    def toggle_level_combo(self, var, combo):
        combo.config(state='readonly' if var.get() else 'disabled')

    def show_markdown_preview(self, text_widget):
        preview_window = tk.Toplevel(self.root)
        preview_window.title("Предпросмотр разметки")
        preview_window.geometry("600x400")

        preview_text = tk.Text(preview_window, wrap='word', font=('Arial', 10))
        preview_text.pack(fill='both', expand=True, padx=10, pady=10)

        text = text_widget.get("1.0", tk.END)
        lines = text.split('\n')

        for line in lines:
            line = line.rstrip()
            if line.startswith('- '):
                preview_text.insert(tk.END, '• ' + line[2:] + '\n')
            elif line.startswith('# '):
                preview_text.insert(tk.END, 'ЗАГОЛОВОК: ' + line[2:] + '\n', 'header')
            elif '**' in line:
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        preview_text.insert(tk.END, part.upper(), 'bold')  # ЗАГЛАВНЫЕ буквы как в требованиях
                    else:
                        preview_text.insert(tk.END, part)
                preview_text.insert(tk.END, '\n')
            elif '[' in line and ']' in line and '(' in line and ')' in line:
                start = line.find('[')
                end = line.find(']')
                link_start = line.find('(')
                link_end = line.find(')')
                link_text = line[start + 1:end]
                link_url = line[link_start + 1:link_end]
                preview_text.insert(tk.END, line[:start])
                preview_text.insert(tk.END, link_text + ' ', 'link')
                preview_text.insert(tk.END, f'({link_url})')
                preview_text.insert(tk.END, line[link_end + 1:] + '\n')
            else:
                preview_text.insert(tk.END, line + '\n')

        preview_text.tag_configure('header', font=('Arial', 12, 'bold'))
        preview_text.tag_configure('bold', font=('Arial', 10, 'bold'))
        preview_text.tag_configure('link', foreground='blue', underline=True)

        preview_text.config(state='disabled')

    def refresh_goals(self):
        for item in self.goals_tree.get_children():
            self.goals_tree.delete(item)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT id, название, тип, статус FROM цели")
        goals = c.fetchall()

        for goal in goals:
            self.goals_tree.insert('', 'end', values=goal)

        conn.close()

    def save_goal(self):
        name = self.name_entry.get()
        goal_type = self.type_combo.get()
        status = self.status_combo.get()
        plan_date = self.plan_date_entry.get()
        fact_date = self.fact_date_entry.get()
        temp = self.temp_entry.get()
        description = self.desc_text.get("1.0", tk.END).strip()

        if not name or not goal_type or not status:
            messagebox.showerror("Ошибка", "Заполните обязательные поля: Название, Тип, Статус")
            return

        skills_text = self.skills_combo.get().strip()
        skills_list = []
        if skills_text:
            skills_list = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
            if len(skills_list) > 3:
                messagebox.showerror("Ошибка", "Можно указать не более 3 навыков")
                return

        selected_competencies = []
        for var, level_combo, comp_id in self.competencies_vars:
            if var.get():
                level = level_combo.get()
                if not level:
                    messagebox.showerror("Ошибка", "Для выбранных компетенций укажите уровень")
                    return
                selected_competencies.append((comp_id, level))

        if len(selected_competencies) > 3:
            messagebox.showerror("Ошибка", "Можно выбрать не более 3 компетенций")
            return

        if len(selected_competencies) < 1:
            messagebox.showerror("Ошибка", "Выберите хотя бы 1 компетенцию")
            return

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()

        c.execute('''
            INSERT INTO цели (название, тип, статус, план_дата, факт_дата, темп, описание)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (name, goal_type, status, plan_date, fact_date, temp, description))

        goal_id = c.lastrowid

        for skill_name in skills_list:
            c.execute("SELECT id FROM навыка WHERE название = ?", (skill_name,))
            skill_row = c.fetchone()

            if skill_row:
                skill_id = skill_row[0]
            else:
                c.execute("INSERT INTO навыка (название) VALUES (?)", (skill_name,))
                skill_id = c.lastrowid

            c.execute("INSERT INTO цель_навыки (цель_id, навык_id) VALUES (?, ?)", (goal_id, skill_id))

        for comp_id, level in selected_competencies:
            c.execute('''
                INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень)
                VALUES (?, ?, ?)
            ''', (goal_id, comp_id, int(level)))

        conn.commit()
        conn.close()

        messagebox.showinfo("Успех", "Цель добавлена")
        self.add_window.destroy()
        self.refresh_goals()
        self.check_all_achievements()
        self.update_stats()
        self.update_semester_progress_auto()

    def edit_goal(self):
        selected = self.goals_tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите цель для редактирования")
            return

        item = self.goals_tree.item(selected[0])
        goal_id = item['values'][0]

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT * FROM цели WHERE id = ?", (goal_id,))
        goal = c.fetchone()

        c.execute('''
            SELECT н.название 
            FROM навыка н
            JOIN цель_навыки цн ON н.id = цн.навык_id
            WHERE цн.цель_id = ?
        ''', (goal_id,))
        goal_skills = [skill[0] for skill in c.fetchall()]

        c.execute('''
            SELECT к.id, к.название, цк.уровень
            FROM компетенции к
            JOIN цель_компетенции цк ON к.id = цк.компетенция_id
            WHERE цк.цель_id = ?
        ''', (goal_id,))
        goal_competencies = c.fetchall()

        conn.close()

        if not goal:
            return

        edit_window = tk.Toplevel(self.root)
        edit_window.title("Редактировать цель")
        edit_window.geometry("800x600")

        row = 0

        ttk.Label(edit_window, text="Название:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        name_entry = ttk.Entry(edit_window, width=50)
        name_entry.insert(0, goal[1])
        name_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(edit_window, text="Тип:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        type_combo = ttk.Combobox(edit_window, values=['Курс', 'Проект', 'Самообразование', 'Семинар', 'Другое'],
                                  width=47)
        type_combo.set(goal[2])
        type_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(edit_window, text="Статус:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        status_combo = ttk.Combobox(edit_window, values=['Новая', 'В процессе', 'Завершена', 'Отменена'], width=47)
        status_combo.set(goal[3])
        status_combo.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(edit_window, text="Плановая дата (ГГГГ-ММ-ДД):").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        plan_date_entry = ttk.Entry(edit_window, width=20)
        plan_date_entry.insert(0, goal[4] if goal[4] else '')
        plan_date_entry.grid(row=row, column=1, padx=5, pady=5, sticky='w')

        ttk.Label(edit_window, text="Фактическая дата (ГГГГ-ММ-ДД):").grid(row=row, column=1, sticky='e', padx=20)
        fact_date_entry = ttk.Entry(edit_window, width=20)
        fact_date_entry.insert(0, goal[5] if goal[5] else '')
        fact_date_entry.grid(row=row, column=2, padx=5, pady=5, sticky='w')
        row += 1

        ttk.Label(edit_window, text="Темп:").grid(row=row, column=0, sticky='w', padx=5, pady=5)
        temp_entry = ttk.Entry(edit_window, width=50)
        temp_entry.insert(0, goal[6] if goal[6] else '')
        temp_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(edit_window, text="Навыки (до 3, через запятую):").grid(row=row, column=0, sticky='w', padx=5, pady=5)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT название FROM навыка")
        existing_skills = [skill[0] for skill in c.fetchall()]
        conn.close()

        skills_entry = ttk.Combobox(edit_window, values=existing_skills, width=50)
        skills_entry.set(', '.join(goal_skills))
        skills_entry.grid(row=row, column=1, padx=5, pady=5, columnspan=2)
        row += 1

        ttk.Label(edit_window, text="Компетенции (выберите 1-3):").grid(row=row, column=0, sticky='nw', padx=5, pady=5)
        comp_frame = ttk.Frame(edit_window)
        comp_frame.grid(row=row, column=1, padx=5, pady=5, columnspan=2, sticky='w')

        competencies_vars = []
        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT id, название FROM компетенции")
        all_comps = c.fetchall()

        goal_comp_dict = {comp[0]: comp[2] for comp in goal_competencies}

        for comp_id, comp_name in all_comps:
            var = tk.BooleanVar(value=(comp_id in goal_comp_dict))
            cb = ttk.Checkbutton(comp_frame, text=comp_name, variable=var)
            cb.pack(anchor='w')

            level_combo = ttk.Combobox(comp_frame, values=['1', '2', '3', '4', '5'], width=5, state='readonly')
            level_combo.set(str(goal_comp_dict.get(comp_id, '3')))
            level_combo.pack(anchor='w', pady=2)
            level_combo.config(state='readonly' if var.get() else 'disabled')

            var.trace('w', lambda *args, v=var, c=level_combo: self.toggle_level_combo(v, c))

            competencies_vars.append((var, level_combo, comp_id))

        conn.close()
        row += 1

        desc_frame = ttk.LabelFrame(edit_window, text="Описание (простая разметка)")
        desc_frame.grid(row=row, column=0, columnspan=3, padx=5, pady=5, sticky='nsew')

        desc_text = tk.Text(desc_frame, width=40, height=8)
        desc_text.insert("1.0", goal[7] if goal[7] else '')
        desc_text.pack(side='left', fill='both', expand=True, padx=5, pady=5)

        preview_btn = ttk.Button(desc_frame, text="Просмотр",
                                 command=lambda: self.show_markdown_preview(desc_text))
        preview_btn.pack(side='left', padx=5)

        edit_window.grid_rowconfigure(row, weight=1)
        edit_window.grid_columnconfigure(1, weight=1)
        row += 1

        def save_changes():
            new_name = name_entry.get()
            new_type = type_combo.get()
            new_status = status_combo.get()
            new_plan_date = plan_date_entry.get()
            new_fact_date = fact_date_entry.get()
            new_temp = temp_entry.get()
            new_description = desc_text.get("1.0", tk.END).strip()

            skills_text = skills_entry.get().strip()
            new_skills_list = []
            if skills_text:
                new_skills_list = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
                if len(new_skills_list) > 3:
                    messagebox.showerror("Ошибка", "Можно указать не более 3 навыков")
                    return

            selected_competencies = []
            for var, level_combo, comp_id in competencies_vars:
                if var.get():
                    level = level_combo.get()
                    if not level:
                        messagebox.showerror("Ошибка", "Для выбранных компетенций укажите уровень")
                        return
                    selected_competencies.append((comp_id, level))

            if len(selected_competencies) > 3:
                messagebox.showerror("Ошибка", "Можно выбрать не более 3 компетенций")
                return

            if len(selected_competencies) < 1:
                messagebox.showerror("Ошибка", "Выберите хотя бы 1 компетенцию")
                return

            conn = sqlite3.connect('iom.db')
            c = conn.cursor()

            c.execute('''
                UPDATE цели 
                SET название = ?, тип = ?, статус = ?, план_дата = ?, факт_дата = ?, темп = ?, описание = ?
                WHERE id = ?
            ''', (new_name, new_type, new_status, new_plan_date, new_fact_date, new_temp, new_description, goal_id))

            c.execute("DELETE FROM цель_навыки WHERE цель_id = ?", (goal_id,))

            for skill_name in new_skills_list:
                c.execute("SELECT id FROM навыка WHERE название = ?", (skill_name,))
                skill_row = c.fetchone()

                if skill_row:
                    skill_id = skill_row[0]
                else:
                    c.execute("INSERT INTO навыка (название) VALUES (?)", (skill_name,))
                    skill_id = c.lastrowid

                c.execute("INSERT INTO цель_навыки (цель_id, навык_id) VALUES (?, ?)", (goal_id, skill_id))

            c.execute("DELETE FROM цель_компетенции WHERE цель_id = ?", (goal_id,))

            for comp_id, level in selected_competencies:
                c.execute('''
                    INSERT INTO цель_компетенции (цель_id, компетенция_id, уровень)
                    VALUES (?, ?, ?)
                ''', (goal_id, comp_id, int(level)))

            conn.commit()
            conn.close()

            messagebox.showinfo("Успех", "Цель обновлена")
            edit_window.destroy()
            self.refresh_goals()
            self.check_all_achievements()
            self.update_stats()
            self.update_semester_progress_auto()

        ttk.Button(edit_window, text="Сохранить", command=save_changes).grid(row=row, column=2, sticky='e', padx=5,
                                                                             pady=10)

    def delete_goal(self):
        selected = self.goals_tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите цель для удаления")
            return

        item = self.goals_tree.item(selected[0])
        goal_id = item['values'][0]

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            conn = sqlite3.connect('iom.db')
            c = conn.cursor()
            c.execute("DELETE FROM цели WHERE id = ?", (goal_id,))
            conn.commit()
            conn.close()

            self.refresh_goals()
            self.check_all_achievements()
            self.update_stats()
            self.update_semester_progress_auto()

    def check_all_achievements(self):
        conn = sqlite3.connect('iom.db')
        c = conn.cursor()

        c.execute("UPDATE достижения SET получено = 0")

        c.execute("SELECT COUNT(*) FROM цели")
        goal_count = c.fetchone()[0]
        if goal_count >= 1:
            c.execute("UPDATE достижения SET получено = 1 WHERE код = 'ach1'")

        c.execute('''
            SELECT COUNT(*) FROM цели 
            WHERE статус = 'Завершена' 
            AND факт_дата IS NOT NULL 
            AND план_дата IS NOT NULL
            AND факт_дата <= план_дата
        ''')
        timely_goals = c.fetchone()[0]
        if timely_goals >= 3:
            c.execute("UPDATE достижения SET получено = 1 WHERE код = 'ach2'")

        c.execute("SELECT COUNT(DISTINCT тип) FROM цели WHERE статус = 'Завершена'")
        distinct_types = c.fetchone()[0]
        if distinct_types >= 3:
            c.execute("UPDATE достижения SET получено = 1 WHERE код = 'ach3'")

        c.execute('''
            SELECT н.название, COUNT(цн.цель_id) 
            FROM навыка н
            JOIN цель_навыки цн ON н.id = цн.навык_id
            JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершена'
            GROUP BY н.id
            HAVING COUNT(цн.цель_id) >= 4
        ''')
        if c.fetchone():
            c.execute("UPDATE достижения SET получено = 1 WHERE код = 'ach4'")

        c.execute("SELECT COUNT(*) FROM цели WHERE статус = 'В процессе'")
        in_progress = c.fetchone()[0]
        if in_progress >= 5:
            c.execute("UPDATE достижения SET получено = 1 WHERE код = 'ach5'")

        conn.commit()
        conn.close()

        self.update_achievements_list()

    def update_stats(self):
        self.update_profile_stats()
        self.update_competencies_stats()

    def create_profile_tab(self):
        title_frame = ttk.Frame(self.tab_profile)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Мой профиль", font=("Arial", 14, "bold")).pack()

        stats_frame = ttk.LabelFrame(self.tab_profile, text="Статистика")
        stats_frame.pack(fill='both', expand=True, padx=10, pady=10)

        skills_frame = ttk.LabelFrame(stats_frame, text="Навыки")
        skills_frame.pack(fill='both', expand=True, padx=5, pady=5)

        self.skills_tree = ttk.Treeview(skills_frame, columns=('Навык', 'Количество целей'), show='headings', height=8)
        self.skills_tree.heading('Навык', text='Навык')
        self.skills_tree.heading('Количество целей', text='Количество целей')
        self.skills_tree.pack(fill='both', expand=True, padx=5, pady=5)

        types_frame = ttk.LabelFrame(stats_frame, text="Статистика по типам целей")
        types_frame.pack(fill='both', expand=True, padx=5, pady=5)

        self.types_tree = ttk.Treeview(types_frame, columns=('Тип', 'Завершено', 'Всего'), show='headings', height=5)
        self.types_tree.heading('Тип', text='Тип')
        self.types_tree.heading('Завершено', text='Завершено')
        self.types_tree.heading('Всего', text='Всего')
        self.types_tree.pack(fill='both', expand=True, padx=5, pady=5)

        timely_frame = ttk.Frame(stats_frame)
        timely_frame.pack(fill='x', padx=5, pady=5)
        self.timely_label = ttk.Label(timely_frame, text="Процент целей, завершённых в срок: 0%")
        self.timely_label.pack()

        ttk.Button(stats_frame, text="Обновить статистику", command=self.update_profile_stats).pack(pady=10)

        self.update_profile_stats()

    def update_profile_stats(self):
        for tree in [self.skills_tree, self.types_tree]:
            for item in tree.get_children():
                tree.delete(item)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()

        c.execute('''
            SELECT н.название, COUNT(цн.цель_id) as количество
            FROM навыка н
            LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
            LEFT JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершена'
            GROUP BY н.id
            HAVING количество > 0
        ''')
        skills = c.fetchall()
        for skill in skills:
            self.skills_tree.insert('', 'end', values=skill)

        c.execute('''
            SELECT тип, 
                   SUM(CASE WHEN статус = 'Завершена' THEN 1 ELSE 0 END) as завершено,
                   COUNT(*) as всего
            FROM цели
            GROUP BY тип
        ''')
        types = c.fetchall()
        for type_stat in types:
            self.types_tree.insert('', 'end', values=type_stat)

        c.execute("SELECT COUNT(*) FROM цели WHERE статус = 'Завершена' AND факт_дата IS NOT NULL")
        completed_total = c.fetchone()[0]

        c.execute('''
            SELECT COUNT(*) FROM цели 
            WHERE статус = 'Завершена' 
            AND факт_дата IS NOT NULL 
            AND план_дата IS NOT NULL
            AND факт_дата <= план_дата
        ''')
        timely_completed = c.fetchone()[0]

        if completed_total > 0:
            percentage = (timely_completed / completed_total) * 100
            self.timely_label.config(text=f"Процент целей, завершённых в срок: {percentage:.1f}%")
        else:
            self.timely_label.config(text="Процент целей, завершённых в срок: 0%")

        conn.close()

    def create_competencies_tab(self):
        title_frame = ttk.Frame(self.tab_competencies)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Компетенции", font=("Arial", 14, "bold")).pack()

        comp_frame = ttk.LabelFrame(self.tab_competencies, text="Компетенции и уровни")
        comp_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.competencies_tree = ttk.Treeview(comp_frame, columns=('Компетенция', 'Категория', 'Средний уровень'),
                                              show='headings', height=10)
        self.competencies_tree.heading('Компетенция', text='Компетенция')
        self.competencies_tree.heading('Категория', text='Категория')
        self.competencies_tree.heading('Средний уровень', text='Средний уровень')
        self.competencies_tree.pack(fill='both', expand=True, padx=5, pady=5)

        weak_frame = ttk.LabelFrame(self.tab_competencies, text="Слабые зоны (уровень < 3)")
        weak_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.weak_zones_text = tk.Text(weak_frame, height=5, width=50)
        self.weak_zones_text.pack(fill='both', expand=True, padx=5, pady=5)

        rec_frame = ttk.LabelFrame(self.tab_competencies, text="Рекомендации")
        rec_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.recommendations_text = tk.Text(rec_frame, height=5, width=50)
        self.recommendations_text.pack(fill='both', expand=True, padx=5, pady=5)

        ttk.Button(self.tab_competencies, text="Обновить компетенции", command=self.update_competencies_stats).pack(
            pady=10)

        self.update_competencies_stats()

    def update_competencies_stats(self):
        for item in self.competencies_tree.get_children():
            self.competencies_tree.delete(item)

        self.weak_zones_text.delete('1.0', tk.END)
        self.recommendations_text.delete('1.0', tk.END)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()

        c.execute('''
            SELECT к.название, к.категория, ROUND(AVG(цк.уровень), 1) as средний_уровень
            FROM компетенции к
            LEFT JOIN цель_компетенции цк ON к.id = цк.компетенция_id
            LEFT JOIN цели ц ON цк.цель_id = ц.id AND ц.статус = 'Завершена'
            GROUP BY к.id
        ''')
        competencies = c.fetchall()

        weak_zones = []
        recommendations = []

        for comp in competencies:
            self.competencies_tree.insert('', 'end', values=comp)

            if comp[2] is not None and comp[2] < 3:
                weak_zones.append(f"{comp[0]} - {comp[2]}")

                if comp[0] == "Презентация результатов":
                    recommendations.append(
                        "Вы почти не развиваете компетенцию 'Презентация результатов'. Рекомендуем выступить на студенческой конференции.")
                elif comp[0] == "Работа с БД":
                    recommendations.append(
                        "Для развития компетенции 'Работа с БД' пройдите курс по SQL или поработайте над проектом с базами данных.")
                elif comp[0] == "Управление проектами":
                    recommendations.append(
                        "Для развития 'Управления проектами' возьмите на себя роль тимлида в учебном проекте.")
                else:
                    recommendations.append(
                        f"Для развития компетенции '{comp[0]}' рекомендуется выполнить практические задания.")

        if weak_zones:
            self.weak_zones_text.insert('1.0', '\n'.join(weak_zones))
        else:
            self.weak_zones_text.insert('1.0', "Слабых зон не обнаружено")

        if recommendations:
            self.recommendations_text.insert('1.0', '\n\n'.join(recommendations))
        else:
            self.recommendations_text.insert('1.0',
                                             "Все компетенции развиваются хорошо. Продолжайте в том же духе!")

        conn.close()

    def create_achievements_tab(self):
        title_frame = ttk.Frame(self.tab_achievements)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Достижения", font=("Arial", 14, "bold")).pack()

        self.achievements_tree = ttk.Treeview(self.tab_achievements,
                                              columns=('Получено', 'Название', 'Описание'), show='headings',
                                              height=10)
        self.achievements_tree.heading('Получено', text='Получено')
        self.achievements_tree.heading('Название', text='Название')
        self.achievements_tree.heading('Описание', text='Описание')

        self.achievements_tree.column('Получено', width=80)
        self.achievements_tree.column('Название', width=150)
        self.achievements_tree.column('Описание', width=400)

        self.achievements_tree.pack(fill='both', expand=True, padx=10, pady=10)

        ttk.Button(self.tab_achievements, text="Обновить достижения",
                   command=self.update_achievements_list).pack(pady=10)

        self.update_achievements_list()

    def update_achievements_list(self):
        for item in self.achievements_tree.get_children():
            self.achievements_tree.delete(item)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute("SELECT * FROM достижения ORDER BY получено DESC, название")
        achievements = c.fetchall()

        for ach in achievements:
            status = "✓ Получено" if ach[3] == 1 else "✗ Не получено"
            self.achievements_tree.insert('', 'end', values=(status, ach[1], ach[2]))

        conn.close()

    def create_semester_tab(self):
        title_frame = ttk.Frame(self.tab_semester)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Цели на семестр", font=("Arial", 14, "bold")).pack()

        goals_frame = ttk.LabelFrame(self.tab_semester, text="Цели семестра")
        goals_frame.pack(fill='both', expand=True, padx=10, pady=10)

        self.semester_tree = ttk.Treeview(goals_frame, columns=('ID', 'Цель', 'Тип', 'Прогресс'),
                                          show='headings', height=10)
        self.semester_tree.heading('ID', text='ID')
        self.semester_tree.heading('Цель', text='Цель')
        self.semester_tree.heading('Тип', text='Тип')
        self.semester_tree.heading('Прогресс', text='Прогресс')
        self.semester_tree.pack(fill='both', expand=True, padx=5, pady=5)

        btn_frame = ttk.Frame(goals_frame)
        btn_frame.pack(fill='x', pady=5)

        ttk.Button(btn_frame, text="Добавить цель", command=self.add_semester_goal).pack(side='left',
                                                                                         padx=2)
        ttk.Button(btn_frame, text="Удалить цель", command=self.delete_semester_goal).pack(side='left',
                                                                                           padx=2)
        ttk.Button(btn_frame, text="Обновить", command=self.update_semester_progress_auto).pack(
            side='left', padx=2)

        ttk.Button(self.tab_semester, text="Сформировать отчёт", command=self.export_to_word).pack(pady=10)

        self.update_semester_progress_auto()

    def add_semester_goal(self):
        add_window = tk.Toplevel(self.root)
        add_window.title("Добавить цель на семестр")

        ttk.Label(add_window, text="Текст цели:").grid(row=0, column=0, sticky='w', padx=5, pady=5)
        goal_entry = ttk.Entry(add_window, width=40)
        goal_entry.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(add_window, text="Тип цели:").grid(row=1, column=0, sticky='w', padx=5, pady=5)
        type_combo = ttk.Combobox(add_window, values=['Количество', 'Повышение компетенции', 'Другое'])
        type_combo.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(add_window, text="Параметр (опционально):").grid(row=2, column=0, sticky='w', padx=5,
                                                                   pady=5)
        param_entry = ttk.Entry(add_window, width=40)
        param_entry.grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(add_window, text="Целевой прогресс:").grid(row=3, column=0, sticky='w', padx=5, pady=5)
        target_spinbox = ttk.Spinbox(add_window, from_=1, to=100, width=10)
        target_spinbox.grid(row=3, column=1, padx=5, pady=5)
        target_spinbox.set(1)

        def save_semester_goal():
            conn = sqlite3.connect('iom.db')
            c = conn.cursor()
            c.execute('''
                INSERT INTO цель_каса (текст_цели, тип_цели, параметр, целевой_прогресс)
                VALUES (?, ?, ?, ?)
            ''', (goal_entry.get(), type_combo.get(), param_entry.get(), int(target_spinbox.get())))
            conn.commit()
            conn.close()

            messagebox.showinfo("Успех", "Цель на семестр добавлена")
            add_window.destroy()
            self.update_semester_progress_auto()

        ttk.Button(add_window, text="Сохранить", command=save_semester_goal).grid(row=4, column=1,
                                                                                  sticky='e', padx=5,
                                                                                  pady=10)

    def update_semester_progress_auto(self):
        conn = sqlite3.connect('iom.db')
        c = conn.cursor()

        c.execute("SELECT id, текст_цели, тип_цели, параметр, целевой_прогресс FROM цель_каса")
        semester_goals = c.fetchall()

        for goal in semester_goals:
            goal_id, text, goal_type, param, target = goal
            current = 0

            if goal_type == 'Количество':
                if 'курс' in text.lower() or 'курс' in param.lower() if param else False:
                    c.execute("SELECT COUNT(*) FROM цели WHERE тип = 'Курс' AND статус = 'Завершена'")
                    current = c.fetchone()[0]
                elif 'проект' in text.lower() or 'проект' in param.lower() if param else False:
                    c.execute("SELECT COUNT(*) FROM цели WHERE тип = 'Проект' AND статус = 'Завершена'")
                    current = c.fetchone()[0]
                elif 'семинар' in text.lower() or 'семинар' in param.lower() if param else False:
                    c.execute("SELECT COUNT(*) FROM цели WHERE тип = 'Семинар' AND статус = 'Завершена'")
                    current = c.fetchone()[0]
                else:
                    c.execute("SELECT COUNT(*) FROM цели WHERE статус = 'Завершена'")
                    current = c.fetchone()[0]

            elif goal_type == 'Повышение компетенции' and param:
                c.execute('''
                    SELECT ROUND(AVG(цк.уровень), 0)
                    FROM компетенции к
                    JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                    JOIN цели ц ON цк.цель_id = ц.id AND ц.статус = 'Завершена'
                    WHERE к.название = ?
                ''', (param,))
                result = c.fetchone()
                current = int(result[0]) if result and result[0] else 0

            c.execute('''
                UPDATE цель_каса 
                SET текущий_прогресс = ?
                WHERE id = ?
            ''', (current, goal_id))

        conn.commit()
        conn.close()

        self.update_semester_goals_list()

    def update_semester_goals_list(self):
        for item in self.semester_tree.get_children():
            self.semester_tree.delete(item)

        conn = sqlite3.connect('iom.db')
        c = conn.cursor()
        c.execute(
            "SELECT id, текст_цели, тип_цели, текущий_прогресс, целевой_прогресс FROM цель_каса")
        goals = c.fetchall()

        for goal in goals:
            progress_text = f"{goal[3]} из {goal[4]}"
            self.semester_tree.insert('', 'end',
                                      values=(goal[0], goal[1], goal[2], progress_text))

        conn.close()

    def delete_semester_goal(self):
        selected = self.semester_tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите цель для удаления")
            return

        item = self.semester_tree.item(selected[0])
        goal_id = item['values'][0]

        if messagebox.askyesno("Подтверждение", "Удалить выбранную цель?"):
            conn = sqlite3.connect('iom.db')
            c = conn.cursor()
            c.execute("DELETE FROM цель_каса WHERE id = ?", (goal_id,))
            conn.commit()
            conn.close()

            self.update_semester_progress_auto()

    def create_settings_tab(self):
        title_frame = ttk.Frame(self.tab_settings)
        title_frame.pack(fill='x', pady=10)
        ttk.Label(title_frame, text="Настройки", font=("Arial", 14, "bold")).pack()

        spec_frame = ttk.LabelFrame(self.tab_settings, text="Специальность")
        spec_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(spec_frame, text="Выберите специальность:").pack(anchor='w', padx=5, pady=5)

        self.specialty_combo = ttk.Combobox(spec_frame,
                                            values=['Информационные системы',
                                                    'Программная инженерия',
                                                    'Прикладная информатика',
                                                    'Другая'])
        self.specialty_combo.pack(fill='x', padx=5, pady=5)
        self.specialty_combo.set('Информационные системы')

        ttk.Button(spec_frame, text="Сохранить специальность",
                   command=self.save_specialty).pack(pady=5)

        db_frame = ttk.LabelFrame(self.tab_settings, text="База данных")
        db_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(db_frame, text="Текущая БД: SQLite (iom.db)").pack(anchor='w', padx=5, pady=5)

        ttk.Button(db_frame, text="Очистить все данные",
                   command=self.clear_all_data).pack(pady=5)

        info_frame = ttk.LabelFrame(self.tab_settings, text="О программе")
        info_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(info_frame, text="Планировщик индивидуального образовательного маршрута").pack(
            anchor='w', padx=5, pady=2)
        ttk.Label(info_frame, text="Версия 1.0").pack(anchor='w', padx=5, pady=2)
        ttk.Label(info_frame, text="Работает автономно, без подключения к интернету").pack(anchor='w',
                                                                                           padx=5,
                                                                                           pady=2)

    def save_specialty(self):
        specialty = self.specialty_combo.get()
        messagebox.showinfo("Сохранено", f"Специальность '{specialty}' сохранена")

        if os.path.exists('competencies.json'):
            conn = sqlite3.connect('iom.db')
            c = conn.cursor()

            c.execute("DELETE FROM компетенции")

            with open('competencies.json', 'r', encoding='utf-8') as f:
                competencies = json.load(f)
                for comp in competencies:
                    c.execute("INSERT INTO компетенции (название, категория) VALUES (?, ?)",
                              (comp['название'], comp['категория']))

            conn.commit()
            conn.close()

            self.update_competencies_stats()

    def clear_all_data(self):
        if messagebox.askyesno("Подтверждение",
                               "Вы уверены, что хотите удалить все данные?\nЭто действие нельзя отменить."):
            conn = sqlite3.connect('iom.db')
            c = conn.cursor()

            tables = ['цели', 'навыка', 'цель_навыки', 'компетенции',
                      'цель_компетенции', 'цель_каса']

            for table in tables:
                c.execute(f"DELETE FROM {table}")

            c.execute("UPDATE достижения SET получено = 0")

            if os.path.exists('competencies.json'):
                with open('competencies.json', 'r', encoding='utf-8') as f:
                    competencies = json.load(f)
                    for comp in competencies:
                        c.execute("INSERT INTO компетенции (название, категория) VALUES (?, ?)",
                                  (comp['название'], comp['категория']))

            conn.commit()
            conn.close()

            messagebox.showinfo("Очищено", "Все данные удалены")

            self.refresh_goals()
            self.update_profile_stats()
            self.update_competencies_stats()
            self.update_achievements_list()
            self.update_semester_progress_auto()

    def export_to_word(self):
        try:
            doc = Document()

            title = doc.add_heading('Индивидуальный образовательный маршрут', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Отчёт сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph()

            doc.add_heading('Цели', level=1)

            conn = sqlite3.connect('iom.db')
            c = conn.cursor()

            c.execute("SELECT * FROM цели ORDER BY статус, план_дата")
            goals = c.fetchall()

            if goals:
                for goal in goals:
                    doc.add_heading(goal[1], level=2)
                    doc.add_paragraph(f"Тип: {goal[2]}")
                    doc.add_paragraph(f"Статус: {goal[3]}")
                    doc.add_paragraph(f"Плановая дата: {goal[4] or 'Не указана'}")
                    doc.add_paragraph(f"Фактическая дата: {goal[5] or 'Не указана'}")
                    if goal[6]:
                        doc.add_paragraph(f"Темп: {goal[6]}")

                    if goal[7]:
                        self.add_formatted_text_to_doc(doc, goal[7])

                    doc.add_paragraph()
            else:
                doc.add_paragraph("Цели не добавлены")

            doc.add_heading('Навыки', level=1)
            c.execute('''
                SELECT н.название, COUNT(цн.цель_id) as количество
                FROM навыка н
                LEFT JOIN цель_навыки цн ON н.id = цн.навык_id
                LEFT JOIN цели ц ON цн.цель_id = ц.id AND ц.статус = 'Завершена'
                GROUP BY н.id
                HAVING количество > 0
            ''')
            skills = c.fetchall()

            if skills:
                for skill in skills:
                    doc.add_paragraph(f"{skill[0]} — {skill[1]} цели", style='List Bullet')
            else:
                doc.add_paragraph("Навыки не указаны")

            doc.add_heading('Компетенции', level=1)

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Компетенция'
            hdr_cells[1].text = 'Категория'
            hdr_cells[2].text = 'Средний уровень'

            c.execute('''
                SELECT к.название, к.категория, ROUND(AVG(цк.уровень), 1) as средний_уровень
                FROM компетенции к
                LEFT JOIN цель_компетенции цк ON к.id = цк.компетенция_id
                LEFT JOIN цели ц ON цк.цель_id = ц.id AND ц.статус = 'Завершена'
                GROUP BY к.id
            ''')
            competencies = c.fetchall()

            for comp in competencies:
                row_cells = table.add_row().cells
                row_cells[0].text = comp[0] or ''
                row_cells[1].text = comp[1] or ''
                row_cells[2].text = str(comp[2]) if comp[2] else 'Нет данных'

            doc.add_heading('Слабые зоны', level=1)
            weak_zones = [comp for comp in competencies if comp[2] is not None and comp[2] < 3]

            if weak_zones:
                for zone in weak_zones:
                    doc.add_paragraph(f"{zone[0]} — уровень {zone[2]}", style='List Bullet')
            else:
                doc.add_paragraph("Слабых зон не обнаружено")

            doc.add_heading('Рекомендации', level=1)

            recommendations = []
            for comp in competencies:
                if comp[2] is not None and comp[2] < 3:
                    if comp[0] == "Презентация результатов":
                        recommendations.append(
                            "Вы почти не развиваете компетенцию 'Презентация результатов'. Рекомендуем выступить на студенческой конференции.")
                    elif comp[0] == "Работа с БД":
                        recommendations.append(
                            "Для развития компетенции 'Работа с БД' пройдите курс по SQL или поработайте над проектом с базами данных.")
                    elif comp[0] == "Управление проектами":
                        recommendations.append(
                            "Для развития 'Управления проектами' возьмите на себя роль тимлида в учебном проекте.")

            if recommendations:
                for rec in recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
            else:
                doc.add_paragraph("Все компетенции развиваются хорошо. Продолжайте в том же духе!")

            doc.add_heading('Достижения', level=1)
            c.execute("SELECT * FROM достижения WHERE получено = 1")
            achievements = c.fetchall()

            if achievements:
                for ach in achievements:
                    doc.add_paragraph(f"{ach[1]} — {ach[2]}", style='List Bullet')
            else:
                doc.add_paragraph("Достижения не получены")

            doc.add_heading('Цели на семестр', level=1)
            c.execute("SELECT текст_цели, текущий_прогресс, целевой_прогресс FROM цель_каса")
            semester_goals = c.fetchall()

            if semester_goals:
                for goal in semester_goals:
                    doc.add_paragraph(f"{goal[0]} — {goal[1]} из {goal[2]}", style='List Bullet')
            else:
                doc.add_paragraph("Цели на семестр не установлены")

            conn.close()

            filename = f"Отчет_ИОМ_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            messagebox.showinfo("Успех", f"Отчёт сохранён в файл: {filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать отчёт: {str(e)}")

    def add_formatted_text_to_doc(self, doc, text):
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            if line.startswith('# '):
                doc.add_heading(line[2:], level=2)
            elif line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
            elif '**' in line:
                parts = line.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            elif '*' in line:
                parts = line.split('*')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        run = p.add_run(part)
                        run.italic = True
                    else:
                        p.add_run(part)
            elif '__' in line:
                parts = line.split('__')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        run = p.add_run(part)
                        run.italic = True
                    else:
                        p.add_run(part)
            else:
                if '[' in line and ']' in line and '(' in line and ')' in line:
                    start = line.find('[')
                    end = line.find(']')
                    link_start = line.find('(')
                    link_end = line.find(')')
                    link_text = line[start + 1:end]
                    link_url = line[link_start + 1:link_end]

                    p = doc.add_paragraph()
                    p.add_run(line[:start])

                    # Создаем гиперссылку
                    from docx.oxml.shared import qn
                    from docx.oxml import OxmlElement

                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), f'rId{doc.part.next_id}')

                    run = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')

                    # Синий цвет
                    color = OxmlElement('w:color')
                    color.set(qn('w:val'), '0000FF')
                    rPr.append(color)

                    # Подчеркивание
                    u = OxmlElement('w:u')
                    u.set(qn('w:val'), 'single')
                    rPr.append(u)

                    run.append(rPr)
                    text_element = OxmlElement('w:t')
                    text_element.text = link_text
                    run.append(text_element)

                    hyperlink.append(run)
                    p._element.append(hyperlink)

                    # Добавляем остальной текст
                    if line[link_end + 1:]:
                        p.add_run(line[link_end + 1:])
                else:
                    doc.add_paragraph(line)


if __name__ == "__main__":
    root = tk.Tk()
    app = IOMApp(root)
    root.mainloop()