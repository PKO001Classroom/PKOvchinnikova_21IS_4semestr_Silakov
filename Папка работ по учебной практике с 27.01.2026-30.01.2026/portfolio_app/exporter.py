# exporter.py
import os
from datetime import datetime


class ReportGenerator:
    def __init__(self, database):
        self.db = database
        self.ensure_folders()

    def ensure_folders(self):
        """Создание необходимых папок"""
        folders = ["reports", "screenshots"]
        for folder in folders:
            if not os.path.exists(folder):
                os.makedirs(folder)

    def generate_excel_report(self):
        """Создание Excel отчета"""
        try:
            # Пробуем импортировать openpyxl
            try:
                from openpyxl import Workbook
                from openpyxl.chart import BarChart, Reference
                from openpyxl.styles import Font

                # Получаем данные
                stats = self.db.get_statistics()
                entries = self.db.get_all_entries()

                # Создаем рабочую книгу
                wb = Workbook()

                # Лист со статистикой
                ws_stats = wb.active
                ws_stats.title = "Статистика"

                # Заголовок
                ws_stats['A1'] = "ОТЧЕТ ПОРТФОЛИО"
                ws_stats['A1'].font = Font(bold=True, size=14)
                ws_stats['A2'] = f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
                ws_stats['A3'] = f"Всего записей: {stats['total']}"
                ws_stats['A4'] = f"Уникальных соавторов: {stats['unique_coauthors']}"

                # Статистика по типам
                ws_stats['A6'] = "ТИП ЗАПИСИ"
                ws_stats['B6'] = "КОЛИЧЕСТВО"
                ws_stats['A6'].font = Font(bold=True)
                ws_stats['B6'].font = Font(bold=True)

                row = 7
                for item in stats['by_type']:
                    ws_stats[f'A{row}'] = item['entry_type']
                    ws_stats[f'B{row}'] = item['count']
                    row += 1

                # Итог
                ws_stats[f'A{row}'] = "ВСЕГО"
                ws_stats[f'B{row}'] = stats['total']
                ws_stats[f'A{row}'].font = Font(bold=True)
                ws_stats[f'B{row}'].font = Font(bold=True)

                # Создаем график
                chart = BarChart()
                chart.title = "Распределение записей по типам"
                chart.x_axis.title = "Тип записи"
                chart.y_axis.title = "Количество"

                data = Reference(ws_stats, min_col=2, min_row=6, max_row=row - 1)
                categories = Reference(ws_stats, min_col=1, min_row=7, max_row=row - 1)

                chart.add_data(data, titles_from_data=True)
                chart.set_categories(categories)

                ws_stats.add_chart(chart, "D6")

                # Лист с записями
                ws_entries = wb.create_sheet("Записи")

                # Заголовки
                headers = ["ID", "Название", "Тип", "Год", "Создано", "Соавторы"]
                for col, header in enumerate(headers, 1):
                    cell = ws_entries.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)

                # Данные
                row = 2
                for entry in entries:
                    ws_entries.cell(row=row, column=1, value=entry['id'])
                    ws_entries.cell(row=row, column=2, value=entry['title'])
                    ws_entries.cell(row=row, column=3, value=entry['entry_type'])
                    ws_entries.cell(row=row, column=4, value=entry['year'])

                    # Форматируем дату
                    created_at = entry['created_at']
                    if created_at:
                        if isinstance(created_at, str):
                            date_str = created_at[:19]
                        else:
                            date_str = created_at.strftime("%d.%m.%Y %H:%M")
                        ws_entries.cell(row=row, column=5, value=date_str)

                    # Соавторы
                    coauthors = self.db.get_coauthors(entry['id'])
                    if coauthors:
                        ws_entries.cell(row=row, column=6, value=", ".join(coauthors))

                    row += 1

                # Автонастройка ширины колонок
                for column in ws_entries.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws_entries.column_dimensions[column_letter].width = adjusted_width

                # Сохраняем файл
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"reports/portfolio_excel_{timestamp}.xlsx"
                wb.save(filename)

                print(f"✅ Excel отчет создан: {filename}")
                return filename

            except ImportError:
                print("⚠️ Библиотека openpyxl не установлена")
                return None

        except Exception as e:
            print(f"❌ Ошибка создания Excel отчета: {e}")
            import traceback
            traceback.print_exc()
            return None

    def generate_word_report(self):
        """Создание Word отчета"""
        try:
            # Пробуем импортировать python-docx
            try:
                from docx import Document
                from docx.shared import Inches, Pt, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_TABLE_ALIGNMENT

                # Получаем данные
                stats = self.db.get_statistics()
                entries = self.db.get_all_entries()

                # Создаем документ
                doc = Document()

                # Титульная страница
                title = doc.add_heading('ЭЛЕКТРОННЫЙ ПОРТФОЛИО\nСТУДЕНТА-ИССЛЕДОВАТЕЛЯ', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Информация о дате
                date_para = doc.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                date_run = date_para.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
                date_run.font.size = Pt(12)

                doc.add_page_break()

                # Раздел 1: Общая информация
                doc.add_heading('1. Общая информация', level=1)

                info_table = doc.add_table(rows=3, cols=2)
                info_table.style = 'LightShading'
                info_table.autofit = True

                # Заполняем таблицу
                info_table.cell(0, 0).text = 'Всего записей в портфолио:'
                info_table.cell(0, 1).text = str(stats['total'])
                info_table.cell(1, 0).text = 'Уникальных соавторов:'
                info_table.cell(1, 1).text = str(stats['unique_coauthors'])
                info_table.cell(2, 0).text = 'Дата формирования отчета:'
                info_table.cell(2, 1).text = datetime.now().strftime("%d.%m.%Y %H:%M")

                # Раздел 2: Статистика по типам
                doc.add_heading('2. Статистика по типам записей', level=1)

                if stats['by_type']:
                    stats_table = doc.add_table(rows=len(stats['by_type']) + 2, cols=2)
                    stats_table.style = 'LightShading'

                    # Заголовки
                    stats_table.cell(0, 0).text = 'Тип записи'
                    stats_table.cell(0, 1).text = 'Количество'

                    # Данные
                    for i, item in enumerate(stats['by_type'], 1):
                        stats_table.cell(i, 0).text = item['entry_type']
                        stats_table.cell(i, 1).text = str(item['count'])

                    # Итог
                    last_row = len(stats['by_type']) + 1
                    stats_table.cell(last_row, 0).text = 'ВСЕГО'
                    stats_table.cell(last_row, 1).text = str(stats['total'])
                else:
                    doc.add_paragraph('Нет данных')

                # Раздел 3: Последние записи
                doc.add_heading('3. Последние записи', level=1)

                if entries:
                    for i, entry in enumerate(entries[:5], 1):  # Только 5 последних
                        # Заголовок записи
                        entry_heading = doc.add_heading(f'{i}. {entry["title"]}', level=2)

                        # Информация о записи
                        info_text = f'Тип: {entry["entry_type"]}'
                        if entry['year']:
                            info_text += f' | Год: {entry["year"]}'

                        doc.add_paragraph(info_text)

                        # Соавторы
                        coauthors = self.db.get_coauthors(entry['id'])
                        if coauthors:
                            doc.add_paragraph(f'Соавторы: {", ".join(coauthors)}')

                        doc.add_paragraph()  # Пустая строка
                else:
                    doc.add_paragraph('Записей нет')

                # Раздел 4: Распределение по годам
                doc.add_heading('4. Распределение по годам', level=1)

                if stats['by_year']:
                    year_table = doc.add_table(rows=len(stats['by_year']) + 1, cols=2)
                    year_table.style = 'LightList'

                    # Заголовки
                    year_table.cell(0, 0).text = 'Год'
                    year_table.cell(0, 1).text = 'Количество записей'

                    # Данные
                    for i, item in enumerate(stats['by_year'], 1):
                        year_table.cell(i, 0).text = str(item['year'])
                        year_table.cell(i, 1).text = str(item['count'])
                else:
                    doc.add_paragraph('Нет данных')

                # Заключение
                doc.add_heading('Заключение', level=1)
                doc.add_paragraph(
                    'Данный отчет сформирован автоматически системой "Электронный портфолио студента-исследователя".')

                # Сохраняем файл
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"reports/portfolio_word_{timestamp}.docx"
                doc.save(filename)

                print(f"✅ Word отчет создан: {filename}")
                return filename

            except ImportError:
                print("⚠️ Библиотека python-docx не установлена")
                return None

        except Exception as e:
            print(f"❌ Ошибка создания Word отчета: {e}")
            import traceback
            traceback.print_exc()
            return None

    def generate_simple_report(self):
        """Создание простого текстового отчета (запасной вариант)"""
        try:
            from datetime import datetime

            # Получаем данные
            stats = self.db.get_statistics()
            entries = self.db.get_all_entries()

            # Создаем отчет
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"reports/portfolio_simple_{timestamp}.txt"

            with open(filename, "w", encoding="utf-8") as f:
                f.write("=" * 60 + "\n")
                f.write("ОТЧЕТ ПОРТФОЛИО (УПРОЩЕННАЯ ВЕРСИЯ)\n")
                f.write("=" * 60 + "\n\n")

                f.write(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
                f.write(f"Всего записей: {stats['total']}\n\n")

                f.write("Статистика по типам:\n")
                f.write("-" * 30 + "\n")
                for item in stats['by_type']:
                    f.write(f"{item['entry_type']}: {item['count']}\n")
                f.write(f"Всего: {stats['total']}\n\n")

                f.write("Список записей:\n")
                f.write("-" * 60 + "\n")
                for entry in entries:
                    f.write(f"- {entry['title']} ({entry['entry_type']}, {entry['year']})\n")

            print(f"✅ Простой отчет создан: {filename}")
            return filename

        except Exception as e:
            print(f"❌ Ошибка создания простого отчета: {e}")
            import traceback
            traceback.print_exc()
            return None


# Для тестирования
if __name__ == "__main__":
    print("Модуль ReportGenerator загружен")
    print("Доступные методы:")
    print("- generate_excel_report()")
    print("- generate_word_report()")
    print("- generate_simple_report()")