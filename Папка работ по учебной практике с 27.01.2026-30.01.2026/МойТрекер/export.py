from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


class ReportExporter:
    def __init__(self, db):
        self.db = db

    def export_to_word(self, filename="отчет.docx"):
        """Экспорт отчета в Word"""
        doc = Document()

        # Настройки стиля
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Заголовок
        title = doc.add_heading('Академический отчет', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Подзаголовок
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('Личный трекер академической школы')
        subtitle_run.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Дата генерации
        date_para = doc.add_paragraph()
        date_para.add_run(f"Сгенерировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_paragraph()

        # Раздел 1: Записи портфолио
        doc.add_heading('1. Записи портфолио', level=1)
        entries = self.db.get_all_entries_with_keywords()

        if entries:
            for entry in entries:
                doc.add_heading(entry[1], level=2)  # название

                info = doc.add_paragraph()
                info_run = info.add_run(f"Тип: {entry[2]}")
                info_run.bold = True
                info.add_run(f" | Дата: {entry[3]}")

                if entry[5]:  # соавторы
                    info.add_run(f" | Соавторы: {entry[5]}")

                if entry[6]:  # ключевые слова
                    kw_para = doc.add_paragraph()
                    kw_run = kw_para.add_run("Ключевые слова: ")
                    kw_run.bold = True
                    kw_para.add_run(entry[6])

                if entry[4]:  # описание
                    desc_para = doc.add_paragraph("Описание:")
                    desc_run = desc_para.runs[0]
                    desc_run.bold = True
                    doc.add_paragraph(entry[4])

                doc.add_paragraph()  # пустая строка
        else:
            doc.add_paragraph("Записей пока нет")

        doc.add_page_break()

        # Раздел 2: Исследовательская карта
        doc.add_heading('2. Исследовательская карта', level=1)

        # Ключевые слова
        doc.add_heading('Ключевые слова', level=2)
        keywords_stats = self.db.get_keywords_statistics()

        if keywords_stats:
            for keyword, count in keywords_stats:
                doc.add_paragraph(f"• {keyword} — {count} записей", style='List Bullet')
        else:
            doc.add_paragraph("Ключевые слова не указаны")

        # Соавторы
        doc.add_heading('Соавторы', level=2)
        authors_stats = self.db.get_authors_statistics()

        if authors_stats:
            for author, count in authors_stats:
                doc.add_paragraph(f"• {author} — {count} работ", style='List Bullet')
        else:
            doc.add_paragraph("Соавторы не указаны")

        doc.add_page_break()

        # Раздел 3: Компетенции
        doc.add_heading('3. Профиль компетенций', level=1)
        comp_stats = self.db.get_competencies_statistics()

        if comp_stats:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'

            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Компетенция'
            hdr_cells[1].text = 'Средний уровень'
            hdr_cells[2].text = 'Количество оценок'

            # Данные
            for comp_name, avg_level, count in comp_stats:
                row_cells = table.add_row().cells
                row_cells[0].text = comp_name
                row_cells[1].text = f"{avg_level:.1f}/5"
                row_cells[2].text = str(count)
        else:
            doc.add_paragraph("Нет данных о компетенциях")

        # Рекомендации
        doc.add_heading('Рекомендации', level=2)
        recommendations = self.db.get_recommendations()

        if recommendations:
            for rec in recommendations:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(rec)
        else:
            doc.add_paragraph("Пока нет рекомендаций. Продолжайте добавлять записи!")

        doc.add_page_break()

        # Раздел 4: Достижения
        doc.add_heading('4. Достижения', level=1)
        achievements = self.db.get_achievements()

        if achievements:
            for ach_id, name, description, obtained, date_obtained in achievements:
                if obtained:
                    p = doc.add_paragraph()
                    p.add_run("✅ ").bold = True
                    p.add_run(f"{name}").bold = True
                    p.add_run(f" ({date_obtained})")
                    doc.add_paragraph(f"   {description}")
                else:
                    p = doc.add_paragraph()
                    p.add_run("◻ ").bold = True
                    p.add_run(f"{name}")
                    doc.add_paragraph(f"   {description}")
        else:
            doc.add_paragraph("Достижения пока не получены")

        # Раздел 5: Цели
        doc.add_heading('5. Цели на семестр', level=1)

        # Сохраняем файл
        doc.save(filename)
        return filename