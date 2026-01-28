import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import os
import webbrowser
from database import db_manager
from file_manager import file_manager
from datetime import datetime


class ResearchPortfolioApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Электронный портфолио студента-исследователя")
        self.root.geometry("1200x700")

        self.current_entry_id = None
        self.current_file_path = None

        self.setup_ui()
        self.load_entries()

    def setup_ui(self):
        # Панель управления
        control_frame = ttk.LabelFrame(self.root, text="Управление", padding=10)
        control_frame.pack(fill=tk.X, padx=10, pady=5)

        # Поля ввода
        ttk.Label(control_frame, text="Название:").grid(row=0, column=0, sticky=tk.W)
        self.title_entry = ttk.Entry(control_frame, width=40)
        self.title_entry.grid(row=0, column=1, padx=5)

        ttk.Label(control_frame, text="Тип записи:").grid(row=0, column=2, sticky=tk.W, padx=10)
        self.type_combo = ttk.Combobox(control_frame,
                                       values=["Публикация", "Конференция", "Грант", "Преподавание", "Достижение"],
                                       width=15,
                                       state="readonly")
        self.type_combo.set("Публикация")
        self.type_combo.grid(row=0, column=3, padx=5)

        ttk.Label(control_frame, text="Год:").grid(row=0, column=4, sticky=tk.W, padx=10)
        self.year_entry = ttk.Entry(control_frame, width=10)
        self.year_entry.grid(row=0, column=5, padx=5)

        # Кнопки управления
        button_frame = ttk.Frame(control_frame)
        button_frame.grid(row=1, column=0, columnspan=6, pady=5)

        ttk.Button(button_frame, text="Создать", command=self.create_entry).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Сохранить", command=self.save_entry).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Удалить", command=self.delete_entry).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Открыть описание", command=self.open_description).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Экспорт в Excel", command=self.export_excel).pack(side=tk.LEFT, padx=2)
        ttk.Button(button_frame, text="Экспорт в Word", command=self.export_word).pack(side=tk.LEFT, padx=2)

        # Список записей (Treeview)
        list_frame = ttk.LabelFrame(self.root, text="Список записей", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        columns = ("ID", "Название", "Тип", "Год", "Дата создания")
        self.tree = ttk.Treeview(list_frame, columns=columns, show="headings", height=15)

        # Настройка колонок
        self.tree.heading("ID", text="ID")
        self.tree.heading("Название", text="Название")
        self.tree.heading("Тип", text="Тип")
        self.tree.heading("Год", text="Год")
        self.tree.heading("Дата создания", text="Дата создания")

        self.tree.column("ID", width=50)
        self.tree.column("Название", width=250)
        self.tree.column("Тип", width=120)
        self.tree.column("Год", width=80)
        self.tree.column("Дата создания", width=150)

        # Полоса прокрутки
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Область редактирования
        edit_frame = ttk.LabelFrame(self.root, text="Редактирование описания (Markdown)", padding=10)
        edit_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        self.text_editor = tk.Text(edit_frame, wrap=tk.WORD, height=10, font=("Arial", 10))
        text_scrollbar = ttk.Scrollbar(edit_frame, orient=tk.VERTICAL, command=self.text_editor.yview)
        self.text_editor.configure(yscrollcommand=text_scrollbar.set)

        self.text_editor.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Подсказки по синтаксису
        hint_frame = ttk.Frame(edit_frame)
        hint_frame.pack(fill=tk.X, pady=5)

        ttk.Label(hint_frame, text="Подсказки:").pack(side=tk.LEFT)
        ttk.Label(hint_frame, text="Цитата: > текст").pack(side=tk.LEFT, padx=10)
        ttk.Label(hint_frame, text="Код: ```код```").pack(side=tk.LEFT, padx=10)
        ttk.Label(hint_frame, text="Ссылка: [текст](url)").pack(side=tk.LEFT, padx=10)

        # Панель соавторов
        coauthor_frame = ttk.LabelFrame(self.root, text="Соавторы", padding=10)
        coauthor_frame.pack(fill=tk.X, padx=10, pady=5)

        input_frame = ttk.Frame(coauthor_frame)
        input_frame.pack(fill=tk.X, pady=5)

        ttk.Label(input_frame, text="Имя соавтора:").pack(side=tk.LEFT)
        self.coauthor_entry = ttk.Entry(input_frame, width=30)
        self.coauthor_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(input_frame, text="Добавить соавтора", command=self.add_coauthor).pack(side=tk.LEFT)

        self.coauthor_listbox = tk.Listbox(coauthor_frame, height=3)
        self.coauthor_listbox.pack(fill=tk.X, pady=5)

        # Вкладка аналитики
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        analytics_frame = ttk.Frame(notebook)
        notebook.add(analytics_frame, text="Аналитика и отчётность")

        ttk.Button(analytics_frame, text="Сформировать отчёт",
                   command=self.generate_report, width=20).pack(pady=20)

        ttk.Label(analytics_frame, text="Отчёты будут сохранены в папке 'reports'",
                  font=("Arial", 10)).pack()

        # Привязка событий
        self.tree.bind('<<TreeviewSelect>>', self.on_entry_select)

    def load_entries(self):
        """Загрузка записей из базы данных"""
        try:
            # Очищаем текущий список
            for item in self.tree.get_children():
                self.tree.delete(item)

            # Получаем записи из БД
            entries = db_manager.get_all_entries()

            for entry in entries:
                self.tree.insert("", tk.END, values=entry)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить записи: {e}")

    def create_entry(self):
        """Создание новой записи"""
        try:
            # Получаем данные из полей ввода
            title = self.title_entry.get().strip()
            entry_type = self.type_combo.get()
            year = self.year_entry.get().strip()

            # Проверяем заполненность полей
            if not title:
                messagebox.showwarning("Внимание", "Введите название записи")
                return

            if not year or not year.isdigit():
                messagebox.showwarning("Внимание", "Введите корректный год")
                return

            year = int(year)

            # Создаем временный путь к файлу (будет обновлен позже)
            temp_file_path = f"temp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"

            # Создаем запись в БД
            entry_id = db_manager.create_entry(title, entry_type, year, temp_file_path)

            if entry_id:
                # Создаем реальный файл Markdown
                actual_file_path = file_manager.create_md_file(entry_id, title, "")

                if actual_file_path:
                    # Обновляем путь в БД
                    db_manager.update_entry_file_path(entry_id, actual_file_path)

                    # Очищаем поля ввода
                    self.title_entry.delete(0, tk.END)
                    self.year_entry.delete(0, tk.END)
                    self.text_editor.delete(1.0, tk.END)
                    self.coauthor_listbox.delete(0, tk.END)

                    # Обновляем список
                    self.load_entries()

                    messagebox.showinfo("Успех", f"Запись '{title}' создана! ID: {entry_id}")
                else:
                    # Откатываем создание записи, если файл не создался
                    db_manager.delete_entry(entry_id)
                    messagebox.showerror("Ошибка", "Не удалось создать файл описания")
            else:
                messagebox.showerror("Ошибка", "Не удалось создать запись в базе данных")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при создании записи: {e}")

    def save_entry(self):
        """Сохранение изменений записи"""
        if not self.current_entry_id:
            messagebox.showwarning("Внимание", "Выберите запись для сохранения")
            return

        try:
            # Получаем данные
            title = self.title_entry.get().strip()
            entry_type = self.type_combo.get()
            year = self.year_entry.get().strip()
            content = self.text_editor.get(1.0, tk.END).strip()

            # Проверяем данные
            if not title:
                messagebox.showwarning("Внимание", "Введите название записи")
                return

            if not year or not year.isdigit():
                messagebox.showwarning("Внимание", "Введите корректный год")
                return

            year = int(year)

            # Обновляем запись в БД
            success = db_manager.update_entry(self.current_entry_id, title, entry_type, year)

            if success and self.current_file_path:
                # Сохраняем содержимое в файл
                file_manager.update_md_file(self.current_file_path, content)

                # Обновляем список
                self.load_entries()

                messagebox.showinfo("Успех", "Изменения сохранены!")
            else:
                messagebox.showerror("Ошибка", "Не удалось сохранить изменения")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении: {e}")

    def delete_entry(self):
        """Удаление записи"""
        if not self.current_entry_id:
            messagebox.showwarning("Внимание", "Выберите запись для удаления")
            return

        # Получаем название записи для сообщения
        selected_item = self.tree.selection()
        if selected_item:
            item = self.tree.item(selected_item[0])
            entry_title = item['values'][1] if len(item['values']) > 1 else "запись"
        else:
            entry_title = "запись"

        if messagebox.askyesno("Подтверждение", f"Удалить запись '{entry_title}'?"):
            try:
                # Удаляем запись из БД
                file_path = db_manager.delete_entry(self.current_entry_id)

                if file_path and os.path.exists(file_path):
                    # Удаляем файл
                    file_manager.delete_md_file(file_path)

                # Очищаем поля
                self.clear_fields()

                # Обновляем список
                self.load_entries()

                messagebox.showinfo("Успех", "Запись удалена!")

            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при удалении: {e}")

    def open_description(self):
        """Открытие файла описания"""
        if not self.current_entry_id:
            messagebox.showwarning("Внимание", "Выберите запись")
            return

        if self.current_file_path and os.path.exists(self.current_file_path):
            try:
                file_manager.open_md_file_external(self.current_file_path)
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть файл: {e}")
        else:
            messagebox.showwarning("Внимание", "Файл описания не найден")

    def add_coauthor(self):
        """Добавление соавтора"""
        if not self.current_entry_id:
            messagebox.showwarning("Внимание", "Выберите запись")
            return

        coauthor_name = self.coauthor_entry.get().strip()

        if not coauthor_name:
            messagebox.showwarning("Внимание", "Введите имя соавтора")
            return

        try:
            # Добавляем соавтора в БД
            success = db_manager.add_coauthor(self.current_entry_id, coauthor_name)

            if success:
                # Добавляем в список на экране
                self.coauthor_listbox.insert(tk.END, coauthor_name)
                self.coauthor_entry.delete(0, tk.END)

                # Обновляем список соавторов из БД
                coauthors = db_manager.get_coauthors_by_entry(self.current_entry_id)
                self.coauthor_listbox.delete(0, tk.END)
                for coauthor in coauthors:
                    self.coauthor_listbox.insert(tk.END, coauthor)
            else:
                messagebox.showerror("Ошибка", "Не удалось добавить соавтора")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при добавлении соавтора: {e}")

    def on_entry_select(self, event):
        """Обработка выбора записи в списке"""
        selected_item = self.tree.selection()

        if selected_item:
            item = self.tree.item(selected_item[0])
            values = item['values']

            if len(values) >= 5:
                self.current_entry_id = values[0]  # ID - первый элемент

                # Загружаем данные записи
                entry = db_manager.get_entry_by_id(self.current_entry_id)

                if entry:
                    # Заполняем поля ввода
                    self.title_entry.delete(0, tk.END)
                    self.title_entry.insert(0, entry[1])  # title

                    self.type_combo.set(entry[2])  # entry_type

                    self.year_entry.delete(0, tk.END)
                    self.year_entry.insert(0, str(entry[3]))  # year

                    self.current_file_path = entry[4]  # file_path

                    # Загружаем содержимое файла
                    if self.current_file_path and os.path.exists(self.current_file_path):
                        content = file_manager.read_md_file(self.current_file_path)
                        self.text_editor.delete(1.0, tk.END)
                        self.text_editor.insert(1.0, content)
                    else:
                        self.text_editor.delete(1.0, tk.END)

                    # Загружаем соавторов
                    self.coauthor_listbox.delete(0, tk.END)
                    coauthors = db_manager.get_coauthors_by_entry(self.current_entry_id)
                    for coauthor in coauthors:
                        self.coauthor_listbox.insert(tk.END, coauthor)

                    # Логируем просмотр
                    db_manager.log_activity(self.current_entry_id, "VIEW")

    def clear_fields(self):
        """Очистка полей ввода"""
        self.current_entry_id = None
        self.current_file_path = None

        self.title_entry.delete(0, tk.END)
        self.year_entry.delete(0, tk.END)
        self.text_editor.delete(1.0, tk.END)
        self.coauthor_entry.delete(0, tk.END)
        self.coauthor_listbox.delete(0, tk.END)

        # Снимаем выделение в дереве
        for item in self.tree.selection():
            self.tree.selection_remove(item)

    def export_excel(self):
        """Экспорт в Excel (заглушка)"""
        messagebox.showinfo("Информация", "Экспорт в Excel будет реализован в модуле export_tools.py")

        # Создаем папку reports, если её нет
        if not os.path.exists("reports"):
            os.makedirs("reports")

        # Простое создание Excel файла для демонстрации
        try:
            from openpyxl import Workbook
            import datetime

            wb = Workbook()
            ws = wb.active
            ws.title = "Портфолио"

            # Заголовки
            ws['A1'] = "ID"
            ws['B1'] = "Название"
            ws['C1'] = "Тип"
            ws['D1'] = "Год"
            ws['E1'] = "Дата создания"

            # Получаем данные
            entries = db_manager.get_all_entries()

            # Заполняем данные
            for i, entry in enumerate(entries, start=2):
                ws[f'A{i}'] = entry[0]
                ws[f'B{i}'] = entry[1]
                ws[f'C{i}'] = entry[2]
                ws[f'D{i}'] = entry[3]
                ws[f'E{i}'] = entry[4]

            # Сохраняем файл
            filename = f"reports/portfolio_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            wb.save(filename)

            messagebox.showinfo("Успех", f"Excel файл создан: {filename}")

        except ImportError:
            messagebox.showwarning("Внимание", "Библиотека openpyxl не установлена")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте в Excel: {e}")

    def export_word(self):
        """Экспорт в Word (заглушка)"""
        messagebox.showinfo("Информация", "Экспорт в Word будет реализован в модуле export_tools.py")

        # Создаем папку reports, если её нет
        if not os.path.exists("reports"):
            os.makedirs("reports")

        # Простое создание Word файла для демонстрации
        try:
            from docx import Document
            from docx.shared import Inches
            import datetime

            document = Document()

            # Заголовок
            document.add_heading('Электронный портфолио студента-исследователя', 0)

            # Текущая дата
            document.add_paragraph(f'Дата формирования: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')
            document.add_paragraph()

            # Получаем данные
            entries = db_manager.get_all_entries()

            # Таблица с записями
            if entries:
                document.add_heading('Список записей', level=1)
                table = document.add_table(rows=1, cols=5)
                table.style = 'LightShading-Accent1'

                # Заголовки таблицы
                header_cells = table.rows[0].cells
                header_cells[0].text = 'ID'
                header_cells[1].text = 'Название'
                header_cells[2].text = 'Тип'
                header_cells[3].text = 'Год'
                header_cells[4].text = 'Дата создания'

                # Данные таблицы
                for entry in entries:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(entry[0])
                    row_cells[1].text = entry[1]
                    row_cells[2].text = entry[2]
                    row_cells[3].text = str(entry[3])
                    row_cells[4].text = entry[4]
            else:
                document.add_paragraph('Записей не найдено.')

            # Сохраняем файл
            filename = f"reports/portfolio_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            document.save(filename)

            messagebox.showinfo("Успех", f"Word файл создан: {filename}")

        except ImportError:
            messagebox.showwarning("Внимание", "Библиотека python-docx не установлена")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте в Word: {e}")

    def generate_report(self):
        """Формирование отчёта"""
        # Создаем папку reports, если её нет
        if not os.path.exists("reports"):
            os.makedirs("reports")

        try:
            # Получаем статистику
            stats = db_manager.get_statistics()

            if not stats:
                messagebox.showwarning("Внимание", "Нет данных для формирования отчёта")
                return

            # Создаем Excel отчет
            self.create_excel_report(stats)

            # Создаем Word отчет
            self.create_word_report(stats)

            messagebox.showinfo("Успех", "Отчёты созданы в папке 'reports'")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при формировании отчёта: {e}")

    def create_excel_report(self, stats):
        """Создание Excel отчета"""
        try:
            from openpyxl import Workbook
            from openpyxl.chart import BarChart, Reference
            import datetime

            wb = Workbook()

            # Лист со статистикой
            ws_stats = wb.active
            ws_stats.title = "Статистика"

            # Заголовок
            ws_stats['A1'] = "Отчёт по портфолио"
            ws_stats['A2'] = f"Дата формирования: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}"

            # Статистика по типам
            ws_stats['A4'] = "Распределение по типам:"
            ws_stats['A5'] = "Тип записи"
            ws_stats['B5'] = "Количество"

            row = 6
            for entry_type, count in stats.get('by_type', {}).items():
                ws_stats[f'A{row}'] = entry_type
                ws_stats[f'B{row}'] = count
                row += 1

            # Статистика по годам
            ws_stats['D4'] = "Распределение по годам:"
            ws_stats['D5'] = "Год"
            ws_stats['E5'] = "Количество"

            row = 6
            for year, count in stats.get('by_year', {}).items():
                ws_stats[f'D{row}'] = year
                ws_stats[f'E{row}'] = count
                row += 1

            # Общая статистика
            ws_stats['G4'] = "Общая статистика:"
            ws_stats['G5'] = "Показатель"
            ws_stats['H5'] = "Значение"

            ws_stats['G6'] = "Уникальных соавторов"
            ws_stats['H6'] = stats.get('unique_coauthors', 0)

            # Лист с графиками
            ws_chart = wb.create_sheet("Графики")

            # График по типам
            if stats.get('by_type'):
                chart1 = BarChart()
                chart1.title = "Распределение записей по типам"
                chart1.x_axis.title = "Тип записи"
                chart1.y_axis.title = "Количество"

                data = Reference(ws_stats, min_col=2, min_row=5, max_row=5 + len(stats['by_type']))
                categories = Reference(ws_stats, min_col=1, min_row=6, max_row=5 + len(stats['by_type']))
                chart1.add_data(data, titles_from_data=True)
                chart1.set_categories(categories)
                ws_chart.add_chart(chart1, "A1")

            # Сохраняем файл
            filename = f"reports/portfolio_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            wb.save(filename)

        except Exception as e:
            print(f"Ошибка при создании Excel отчета: {e}")
            raise

    def create_word_report(self, stats):
        """Создание Word отчета"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import datetime

            document = Document()

            # Настройка стилей
            style = document.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Титульный лист
            title = document.add_heading('Электронный портфолио студента-исследователя', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph()
            subtitle = document.add_paragraph('Аналитический отчёт')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_paragraph()
            date_para = document.add_paragraph(
                f'Дата формирования: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            document.add_page_break()

            # Сводная таблица ключевых показателей
            document.add_heading('Ключевые показатели', level=1)

            table = document.add_table(rows=4, cols=2)
            table.style = 'LightShading-Accent1'

            # Заполняем таблицу
            table.cell(0, 0).text = 'Показатель'
            table.cell(0, 1).text = 'Значение'

            table.cell(1, 0).text = 'Всего записей'
            total_entries = sum(stats.get('by_type', {}).values())
            table.cell(1, 1).text = str(total_entries)

            table.cell(2, 0).text = 'Уникальных соавторов'
            table.cell(2, 1).text = str(stats.get('unique_coauthors', 0))

            table.cell(3, 0).text = 'Типы записей'
            table.cell(3, 1).text = ', '.join(stats.get('by_type', {}).keys())

            document.add_paragraph()

            # Распределение по типам
            if stats.get('by_type'):
                document.add_heading('Распределение записей по типам', level=2)
                for entry_type, count in stats['by_type'].items():
                    document.add_paragraph(f'{entry_type}: {count} записей', style='ListBullet')

            document.add_paragraph()

            # Последние 5 записей
            if stats.get('last_5_entries'):
                document.add_heading('Последние записи', level=2)
                for entry in stats['last_5_entries']:
                    if len(entry) >= 3:
                        document.add_paragraph(f'{entry[0]} ({entry[1]}, {entry[2]} г.)', style='ListBullet')

            # Сохраняем файл
            filename = f"reports/portfolio_report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            document.save(filename)

        except Exception as e:
            print(f"Ошибка при создании Word отчета: {e}")
            raise


if __name__ == "__main__":
    # Создаем необходимые папки
    if not os.path.exists("portfolio_md"):
        os.makedirs("portfolio_md")
    if not os.path.exists("reports"):
        os.makedirs("reports")

    try:
        root = tk.Tk()
        app = ResearchPortfolioApp(root)
        root.mainloop()
    except Exception as e:
        print(f"Ошибка запуска приложения: {e}")
        import traceback

        traceback.print_exc()
        messagebox.showerror("Ошибка", f"Не удалось запустить приложение: {e}")