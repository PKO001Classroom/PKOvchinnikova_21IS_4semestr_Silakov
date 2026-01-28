# export_tools.py
import os
from datetime import datetime
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
import pandas as pd
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference, LineChart
from openpyxl.drawing.image import Image
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import matplotlib

matplotlib.use('Agg')  # Для работы без GUI


class ExportTools:
    def __init__(self, db_manager):
        self.db_manager = db_manager
        self.reports_dir = "reports"
        self.ensure_reports_dir()

    def ensure_reports_dir(self):
        """Создание папки для отчетов"""
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir)

    def generate_excel_report(self):
        """Генерация Excel отчета с таблицами и диаграммами"""
        try:
            # Получаем статистику
            stats = self.db_manager.get_statistics()

            # Создаем Excel файл
            wb = Workbook()

            # Лист 1: Статистика
            ws_stats = wb.active
            ws_stats.title = "Статистика"

            # Заголовок
            ws_stats['A1'] = "АНАЛИТИЧЕСКИЙ ОТЧЕТ"
            ws_stats['A1'].font = {'bold': True, 'size': 14}
            ws_stats['A2'] = f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            ws_stats['A3'] = "Электронный портфолио студента-исследователя"

            # Раздел 1: Общая статистика
            ws_stats['A5'] = "1. ОБЩАЯ СТАТИСТИКА"
            ws_stats['A5'].font = {'bold': True}

            general_stats = [
                ["Показатель", "Значение"],
                ["Всего записей", sum(stats.get('by_type', {}).values())],
                ["Уникальных соавторов", stats.get('unique_coauthors', 0)],
                ["Период охвата",
                 f"{min(stats.get('by_year', {}).keys(), default='-')} - {max(stats.get('by_year', {}).keys(), default='-')}"],
                ["Дата первой записи", "Извлекается из БД"],
                ["Дата последней записи", "Извлекается из БД"]
            ]

            for i, row in enumerate(general_stats):
                for j, value in enumerate(row):
                    ws_stats.cell(row=6 + i, column=1 + j, value=value)

            # Раздел 2: Распределение по типам
            ws_stats['A12'] = "2. РАСПРЕДЕЛЕНИЕ ПО ТИПАМ"
            ws_stats['A12'].font = {'bold': True}

            ws_stats['A13'] = "Тип записи"
            ws_stats['B13'] = "Количество"
            ws_stats['A13'].font = {'bold': True}
            ws_stats['B13'].font = {'bold': True}

            row = 14
            for entry_type, count in stats.get('by_type', {}).items():
                ws_stats.cell(row=row, column=1, value=entry_type)
                ws_stats.cell(row=row, column=2, value=count)
                row += 1

            # Раздел 3: Динамика по годам
            ws_stats['D12'] = "3. ДИНАМИКА ПО ГОДАМ"
            ws_stats['D12'].font = {'bold': True}

            ws_stats['D13'] = "Год"
            ws_stats['E13'] = "Количество"
            ws_stats['D13'].font = {'bold': True}
            ws_stats['E13'].font = {'bold': True}

            row = 14
            for year, count in sorted(stats.get('by_year', {}).items()):
                ws_stats.cell(row=row, column=4, value=year)
                ws_stats.cell(row=row, column=5, value=count)
                row += 1

            # Лист 2: Графики
            ws_chart = wb.create_sheet("Графики")

            # Создаем и сохраняем графики
            chart_files = self.create_charts(stats)

            # Вставляем график распределения по типам
            if 'type_chart.png' in chart_files:
                img = Image(chart_files['type_chart.png'])
                img.width = 400
                img.height = 300
                ws_chart.add_image(img, 'A1')

            # Вставляем график динамики по годам
            if 'year_chart.png' in chart_files:
                img = Image(chart_files['year_chart.png'])
                img.width = 400
                img.height = 300
                ws_chart.add_image(img, 'I1')

            # Вставляем график активности
            if 'activity_chart.png' in chart_files:
                img = Image(chart_files['activity_chart.png'])
                img.width = 400
                img.height = 300
                ws_chart.add_image(img, 'A20')

            # Настройка ширины колонок
            for column in ['A', 'B', 'D', 'E']:
                ws_stats.column_dimensions[column].width = 20

            # Сохраняем файл
            filename = os.path.join(self.reports_dir,
                                    f"portfolio_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            wb.save(filename)

            # Удаляем временные файлы графиков
            for chart_file in chart_files.values():
                if os.path.exists(chart_file):
                    os.remove(chart_file)

            return filename

        except Exception as e:
            print(f"Ошибка при создании Excel отчета: {e}")
            return None

    def create_charts(self, stats):
        """Создание графиков и сохранение их как PNG"""
        chart_files = {}

        try:
            # График 1: Распределение по типам
            if stats.get('by_type'):
                plt.figure(figsize=(10, 6))
                types = list(stats['by_type'].keys())
                counts = list(stats['by_type'].values())

                bars = plt.bar(types, counts, color=['#4CAF50', '#2196F3', '#FF9800', '#9C27B0', '#F44336'])
                plt.title('Распределение записей по типам', fontsize=14, fontweight='bold')
                plt.xlabel('Тип записи', fontsize=12)
                plt.ylabel('Количество', fontsize=12)
                plt.xticks(rotation=45)

                # Добавляем значения на столбцы
                for bar, count in zip(bars, counts):
                    plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                             str(count), ha='center', va='bottom')

                plt.tight_layout()
                type_chart_path = os.path.join(self.reports_dir, 'type_chart.png')
                plt.savefig(type_chart_path, dpi=300, bbox_inches='tight')
                chart_files['type_chart.png'] = type_chart_path
                plt.close()

            # График 2: Динамика по годам
            if stats.get('by_year'):
                plt.figure(figsize=(12, 6))
                years = sorted(stats['by_year'].keys())
                counts = [stats['by_year'][year] for year in years]

                plt.plot(years, counts, marker='o', linestyle='-', linewidth=2,
                         markersize=8, color='#2196F3')
                plt.fill_between(years, counts, alpha=0.3, color='#2196F3')

                plt.title('Динамика публикаций по годам', fontsize=14, fontweight='bold')
                plt.xlabel('Год', fontsize=12)
                plt.ylabel('Количество записей', fontsize=12)
                plt.grid(True, alpha=0.3)
                plt.gca().xaxis.set_major_locator(MaxNLocator(integer=True))

                # Добавляем аннотации
                for year, count in zip(years, counts):
                    plt.annotate(str(count), xy=(year, count),
                                 xytext=(0, 10), textcoords='offset points',
                                 ha='center', fontsize=10)

                plt.tight_layout()
                year_chart_path = os.path.join(self.reports_dir, 'year_chart.png')
                plt.savefig(year_chart_path, dpi=300, bbox_inches='tight')
                chart_files['year_chart.png'] = year_chart_path
                plt.close()

            # График 3: Активность за 12 месяцев
            if stats.get('activity_last_12_months'):
                plt.figure(figsize=(12, 6))

                months = []
                activities = []
                for month_data in stats['activity_last_12_months']:
                    if len(month_data) >= 2:
                        month = month_data[0].strftime('%b %Y')
                        months.append(month)
                        activities.append(month_data[1])

                if months and activities:
                    bars = plt.bar(months, activities, color='#FF9800')
                    plt.title('Активность за последние 12 месяцев', fontsize=14, fontweight='bold')
                    plt.xlabel('Месяц', fontsize=12)
                    plt.ylabel('Количество действий', fontsize=12)
                    plt.xticks(rotation=45)

                    # Добавляем значения
                    for bar, activity in zip(bars, activities):
                        plt.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                                 str(activity), ha='center', va='bottom', fontsize=9)

                    plt.tight_layout()
                    activity_chart_path = os.path.join(self.reports_dir, 'activity_chart.png')
                    plt.savefig(activity_chart_path, dpi=300, bbox_inches='tight')
                    chart_files['activity_chart.png'] = activity_chart_path
                    plt.close()

            return chart_files

        except Exception as e:
            print(f"Ошибка при создании графиков: {e}")
            return chart_files

    def generate_word_report(self):
        """Генерация профессионального Word-отчета"""
        try:
            # Получаем статистику
            stats = self.db_manager.get_statistics()

            # Создаем документ
            doc = Document()

            # Настройка страницы
            section = doc.sections[0]
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21.0)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)

            # Настройка стилей
            self.setup_styles(doc)

            # Титульный лист
            self.add_title_page(doc)
            doc.add_page_break()

            # Содержание
            self.add_table_of_contents(doc)
            doc.add_page_break()

            # 1. Введение
            self.add_introduction(doc)

            # 2. Ключевые показатели
            self.add_key_metrics(doc, stats)

            # 3. Детальный анализ
            self.add_detailed_analysis(doc, stats)

            # 4. Графики и визуализация
            self.add_charts_section(doc, stats)

            # 5. Последние записи
            self.add_recent_entries(doc, stats)

            # 6. Заключение
            self.add_conclusion(doc)

            # Сохраняем документ
            filename = os.path.join(self.reports_dir,
                                    f"portfolio_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            doc.save(filename)

            return filename

        except Exception as e:
            print(f"Ошибка при создании Word отчета: {e}")
            return None

    def setup_styles(self, doc):
        """Настройка стилей документа"""
        # Основной стиль
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(6)

        # Заголовок 1
        heading1 = doc.styles['Heading 1']
        heading1.font.name = 'Times New Roman'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.paragraph_format.space_before = Pt(18)
        heading1.paragraph_format.space_after = Pt(12)
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Заголовок 2
        heading2 = doc.styles['Heading 2']
        heading2.font.name = 'Times New Roman'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.paragraph_format.space_before = Pt(12)
        heading2.paragraph_format.space_after = Pt(6)

        # Заголовок 3
        heading3 = doc.styles['Heading 3']
        heading3.font.name = 'Times New Roman'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.paragraph_format.space_before = Pt(6)
        heading3.paragraph_format.space_after = Pt(3)

    def add_title_page(self, doc):
        """Добавление титульного листа"""
        # Название учреждения (если нужно)
        uni_para = doc.add_paragraph()
        uni_run = uni_para.add_run("Министерство науки и высшего образования РФ")
        uni_run.font.name = 'Times New Roman'
        uni_run.font.size = Pt(12)
        uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустая строка
        doc.add_paragraph()

        # Основной заголовок
        title = doc.add_heading('ЭЛЕКТРОННЫЙ ПОРТФОЛИО\nСТУДЕНТА-ИССЛЕДОВАТЕЛЯ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Подзаголовок
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run('Аналитический отчет')
        subtitle_run.font.name = 'Times New Roman'
        subtitle_run.font.size = Pt(14)
        subtitle_run.bold = True
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустые строки
        for _ in range(5):
            doc.add_paragraph()

        # Автор
        author = doc.add_paragraph()
        author_run = author.add_run('Подготовил: студент-исследователь')
        author_run.font.name = 'Times New Roman'
        author_run.font.size = Pt(12)
        author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Дата
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y г.")}')
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Город и год
        city_para = doc.add_paragraph()
        city_run = city_para.add_run('Москва, 2026')
        city_run.font.name = 'Times New Roman'
        city_run.font.size = Pt(12)
        city_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table_of_contents(self, doc):
        """Добавление содержания"""
        toc = doc.add_heading('СОДЕРЖАНИЕ', 1)

        contents = [
            "1. ВВЕДЕНИЕ",
            "2. КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ",
            "3. ДЕТАЛЬНЫЙ АНАЛИЗ",
            "  3.1. Распределение по типам деятельности",
            "  3.2. Динамика по годам",
            "  3.3. Активность пользователя",
            "4. ГРАФИКИ И ВИЗУАЛИЗАЦИЯ",
            "5. ПОСЛЕДНИЕ ДОСТИЖЕНИЯ",
            "6. ЗАКЛЮЧЕНИЕ"
        ]

        for item in contents:
            para = doc.add_paragraph(item)
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)

    def add_introduction(self, doc):
        """Добавление введения"""
        doc.add_heading('1. ВВЕДЕНИЕ', 1)

        intro_text = """
Настоящий отчет представляет собой комплексный анализ научно-исследовательской деятельности, 
зафиксированной в электронном портфолио. Отчет сформирован автоматически на основе данных, 
внесенных в систему учета академических достижений.

Цель отчета — предоставить структурированную информацию о результатах исследовательской работы, 
динамике публикационной активности, участии в конференциях, полученных грантах и других формах 
научной деятельности.

Отчет подготовлен с использованием кроссплатформенного приложения «Электронный портфолио 
студента-исследователя», обеспечивающего интеграцию с системой управления базами данных 
PostgreSQL и автоматическую генерацию аналитических материалов.
        """

        para = doc.add_paragraph(intro_text.strip())
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_key_metrics(self, doc, stats):
        """Добавление ключевых показателей"""
        doc.add_heading('2. КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ', 1)

        total_entries = sum(stats.get('by_type', {}).values())
        unique_coauthors = stats.get('unique_coauthors', 0)
        years = list(stats.get('by_year', {}).keys())

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Shading Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заголовок таблицы
        table.cell(0, 0).text = "Показатель"
        table.cell(0, 1).text = "Значение"

        # Заполнение данных
        metrics = [
            ["Общее количество записей", str(total_entries)],
            ["Уникальных соавторов", str(unique_coauthors)],
            ["Количество лет активности", str(len(years)) if years else "0"],
            ["Период активности", f"{min(years) if years else '-'} - {max(years) if years else '-'}"],
            ["Среднее в год", f"{total_entries / len(years):.1f}" if years and len(years) > 0 else "0.0"]
        ]

        for i, (metric, value) in enumerate(metrics, 1):
            table.cell(i, 0).text = metric
            table.cell(i, 1).text = value

        # Настройка шрифта в таблице
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']

    def add_detailed_analysis(self, doc, stats):
        """Добавление детального анализа"""
        doc.add_heading('3. ДЕТАЛЬНЫЙ АНАЛИЗ', 1)

        # 3.1 Распределение по типам
        doc.add_heading('3.1. Распределение по типам деятельности', 2)

        if stats.get('by_type'):
            for entry_type, count in stats['by_type'].items():
                para = doc.add_paragraph(f"• {entry_type}: {count} записей")
                para.style = doc.styles['List Bullet']
        else:
            doc.add_paragraph("Данные о типах записей отсутствуют.")

        # 3.2 Динамика по годам
        doc.add_heading('3.2. Динамика по годам', 2)

        if stats.get('by_year'):
            table = doc.add_table(rows=len(stats['by_year']) + 1, cols=2)
            table.style = 'Light Shading Accent 2'

            table.cell(0, 0).text = "Год"
            table.cell(0, 1).text = "Количество записей"

            for i, (year, count) in enumerate(sorted(stats['by_year'].items()), 1):
                table.cell(i, 0).text = str(year)
                table.cell(i, 1).text = str(count)
        else:
            doc.add_paragraph("Данные по годам отсутствуют.")

        # 3.3 Активность пользователя
        doc.add_heading('3.3. Активность пользователя', 2)

        if stats.get('activity_last_12_months'):
            activity_text = "За последние 12 месяцев зафиксирована следующая активность:\n"
            for month_data in stats['activity_last_12_months']:
                if len(month_data) >= 2:
                    month = month_data[0].strftime('%B %Y')
                    activity_text += f"• {month}: {month_data[1]} действий\n"

            doc.add_paragraph(activity_text)
        else:
            doc.add_paragraph("Данные об активности отсутствуют.")

    def add_charts_section(self, doc, stats):
        """Добавление раздела с графиками"""
        doc.add_heading('4. ГРАФИКИ И ВИЗУАЛИЗАЦИЯ', 1)

        doc.add_paragraph("""
Ниже представлены графические визуализации данных портфолио, 
позволяющие наглядно оценить динамику и структуру научной деятельности.
        """)

        # Создаем и вставляем графики
        chart_files = self.create_charts(stats)

        if 'type_chart.png' in chart_files and os.path.exists(chart_files['type_chart.png']):
            doc.add_heading('Распределение записей по типам', 3)
            doc.add_picture(chart_files['type_chart.png'], width=Inches(6))
            doc.add_paragraph("Рис. 1. Структура научной деятельности по типам записей")

        if 'year_chart.png' in chart_files and os.path.exists(chart_files['year_chart.png']):
            doc.add_heading('Динамика публикаций по годам', 3)
            doc.add_picture(chart_files['year_chart.png'], width=Inches(6))
            doc.add_paragraph("Рис. 2. Динамика научной продуктивности по годам")

        if 'activity_chart.png' in chart_files and os.path.exists(chart_files['activity_chart.png']):
            doc.add_heading('Активность работы с портфолио', 3)
            doc.add_picture(chart_files['activity_chart.png'], width=Inches(6))
            doc.add_paragraph("Рис. 3. Активность пользователя за последние 12 месяцев")

        # Удаляем временные файлы
        for chart_file in chart_files.values():
            if os.path.exists(chart_file):
                os.remove(chart_file)

    def add_recent_entries(self, doc, stats):
        """Добавление информации о последних записях"""
        doc.add_heading('5. ПОСЛЕДНИЕ ДОСТИЖЕНИЯ', 1)

        doc.add_paragraph("""
В данном разделе представлены последние записи, внесенные в портфолио, 
что позволяет оценить текущую научную активность.
        """)

        if stats.get('last_5_entries'):
            table = doc.add_table(rows=len(stats['last_5_entries']) + 1, cols=3)
            table.style = 'Light Shading Accent 3'

            table.cell(0, 0).text = "Название"
            table.cell(0, 1).text = "Тип"
            table.cell(0, 2).text = "Год"

            for i, entry in enumerate(stats['last_5_entries'], 1):
                if len(entry) >= 3:
                    table.cell(i, 0).text = entry[0]
                    table.cell(i, 1).text = entry[1]
                    table.cell(i, 2).text = str(entry[2])
        else:
            doc.add_paragraph("Последние записи отсутствуют.")

    def add_conclusion(self, doc):
        """Добавление заключения"""
        doc.add_heading('6. ЗАКЛЮЧЕНИЕ', 1)

        conclusion_text = """
Настоящий аналитический отчет демонстрирует системный подход к документированию 
и анализу научно-исследовательской деятельности. Представленные данные позволяют 
оценить продуктивность, динамику развития и направления научной работы.

Использование электронного портфолио способствует структурированию академических 
достижений, обеспечивает удобный доступ к информации и позволяет генерировать 
профессиональные отчеты для различных целей (отчетность, подача заявок на гранты, 
формирование резюме).

Рекомендуется продолжать регулярное ведение портфолио для накопления статистически 
значимых данных и отслеживания долгосрочных тенденций в научной деятельности.
        """

        para = doc.add_paragraph(conclusion_text.strip())
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY